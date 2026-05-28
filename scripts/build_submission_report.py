"""Build the college submission DOCX for the password analyzer project."""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from password_ai import PasswordStrengthAnalyzer, generate_secure_password


OUTPUT = Path(r"C:\Users\Sachin Nishad\Desktop\AI_Powered_Password_Strength_Analyzer_Project_Report.docx")
ASSETS = ROOT / "report_assets"

TITLE = "AI-Powered Password Strength Analyzer"
STUDENT = "Sachin Nishad"
ENROLLMENT = "A9920122000732"
UNIVERSITY = "Amity University Online, Noida, Uttar Pradesh, India"
DEGREE = "Bachelor of Computer Applications"
SESSION = "July 2023"
DATE_TEXT = "May 2026"


def main() -> None:
    doc = Document()
    setup_document(doc)
    build_front_matter(doc)
    build_main_report(doc)
    add_references(doc)
    add_appendices(doc)
    doc.save(OUTPUT)
    print(f"Saved {OUTPUT}")


def setup_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 2
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size in [
        ("Title", 20),
        ("Heading 1", 16),
        ("Heading 2", 14),
        ("Heading 3", 12),
    ]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        if style_name.startswith("Heading"):
            style.font.color.rgb = RGBColor(31, 78, 121)
            style.paragraph_format.space_before = Pt(12)
            style.paragraph_format.space_after = Pt(6)
            style.paragraph_format.page_break_before = False

    code_style = styles.add_style("Code Block", 1)
    code_style.font.name = "Courier New"
    code_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
    code_style.font.size = Pt(9)
    code_style.paragraph_format.line_spacing = 1
    code_style.paragraph_format.space_after = Pt(3)

    add_running_header(doc)


def add_running_header(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = "Running head: AI PASSWORD STRENGTH ANALYZER"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.runs[0].font.name = "Times New Roman"
    paragraph.runs[0].font.size = Pt(9)
    paragraph.runs[0].font.color.rgb = RGBColor(90, 90, 90)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Page ")
    run.font.size = Pt(9)
    add_page_number_field(p)


def add_page_number_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def add_centered(doc: Document, text: str, size: int = 12, bold: bool = False) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)


def add_para(doc: Document, text: str = "", style: str | None = None, align=None) -> None:
    p = doc.add_paragraph(style=style)
    p.paragraph_format.line_spacing = 2
    p.paragraph_format.space_after = Pt(6)
    if align is not None:
        p.alignment = align
    p.add_run(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing = 2
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.line_spacing = 2
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)


def page_break(doc: Document) -> None:
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], caption: str | None = None) -> None:
    if caption:
        add_caption(doc, caption)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shade_cell(cell, "D9EAF7")
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.runs[0].bold = True
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = value
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.line_spacing = 1.2
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def build_front_matter(doc: Document) -> None:
    add_centered(doc, UNIVERSITY, 12, True)
    add_centered(doc, "In partial fulfillment of the requirements for the award of the degree", 12)
    add_centered(doc, DEGREE, 14, True)
    doc.add_paragraph()
    add_centered(doc, "Project Report On", 14, True)
    add_centered(doc, f'"{TITLE}"', 20, True)
    doc.add_paragraph()
    add_centered(doc, f"Submitted by: {STUDENT}", 12, True)
    add_centered(doc, f"Enrollment No.: {ENROLLMENT}", 12, True)
    add_centered(doc, f"Academic Session: {SESSION}", 12)
    add_centered(doc, f"Submission: {DATE_TEXT}", 12)
    doc.add_paragraph()
    add_centered(doc, "Project Guide: ______________________________", 12)
    add_centered(doc, "Designation: ________________________________", 12)
    page_break(doc)

    doc.add_heading("DECLARATION", level=1)
    add_para(
        doc,
        f"I, {STUDENT}, student of the Academic Session {SESSION}, hereby declare that the project titled "
        f'"{TITLE}" submitted to {UNIVERSITY} in partial fulfillment of the requirement for the award of the '
        f"degree {DEGREE}, is an original work prepared by me. The project has not previously formed the basis "
        "for the award of any degree, diploma, certificate, fellowship, or other similar title or recognition.",
    )
    add_para(
        doc,
        "I further declare that the contents of this report are based on my own project development, analysis, "
        "testing, and interpretation. All external references used for academic background and technical support "
        "have been cited in the bibliography section according to the required academic format.",
    )
    doc.add_paragraph("\nName and Signature of the Student")
    doc.add_paragraph(STUDENT)
    doc.add_paragraph("Date: ____________________")
    page_break(doc)

    doc.add_heading("CERTIFICATE", level=1)
    add_para(
        doc,
        f"This is to certify that {STUDENT} of {UNIVERSITY} has carried out the project work presented in this "
        f"project report entitled \"{TITLE}\" for the award of {DEGREE} under my guidance. The project report "
        "embodies results of original work and studies carried out by the student.",
    )
    add_para(
        doc,
        "Certified further, that to the best of my knowledge, the work reported herein does not form the basis "
        "for the award of any other degree to the candidate or to anybody else from this or any other "
        "University or Institution.",
    )
    doc.add_paragraph("\nSignature: ______________________________")
    doc.add_paragraph("Name of Guide: __________________________")
    doc.add_paragraph("Designation: ____________________________")
    doc.add_paragraph("Date: ____________________")
    page_break(doc)

    doc.add_heading("PLAGIARISM REPORT DECLARATION", level=1)
    add_para(
        doc,
        f"I, {STUDENT}, enrolled in the degree program {DEGREE}, certify that the submitted project report titled "
        f'"{TITLE}" has been prepared with original wording, original implementation, and appropriate references. '
        "The official plagiarism report should be attached with the final submission after checking this document "
        "through the university-approved plagiarism tool.",
    )
    add_para(
        doc,
        "Plagiarism percentage after final checking: ________%. The final percentage must remain within the "
        "permissible limit prescribed by the University before submission.",
    )
    doc.add_paragraph("\nName and Signature of the Student")
    doc.add_paragraph(STUDENT)
    page_break(doc)

    doc.add_heading("ABSTRACT", level=1)
    for para in abstract_paragraphs():
        add_para(doc, para)
    page_break(doc)

    doc.add_heading("EXTENDED ABSTRACT", level=1)
    for heading, paragraphs in extended_abstract_sections():
        doc.add_heading(heading, level=2)
        for para in paragraphs:
            add_para(doc, para)
    page_break(doc)

    doc.add_heading("TABLE OF CONTENTS", level=1)
    toc_rows = [
        ["Declaration", "2"],
        ["Certificate", "3"],
        ["Plagiarism Report Declaration", "4"],
        ["Abstract", "5"],
        ["Extended Abstract", "To be updated after final pagination"],
        ["List of Tables", "To be updated after final pagination"],
        ["List of Figures", "To be updated after final pagination"],
        ["Chapter 1: Introduction to the Topic", "To be updated after final pagination"],
        ["Chapter 2: Review of Literature", "To be updated after final pagination"],
        ["Chapter 3: Research Objectives and Methodology", "To be updated after final pagination"],
        ["Chapter 4: Data Analysis, Results, and Interpretation", "To be updated after final pagination"],
        ["Chapter 5: Findings and Conclusion", "To be updated after final pagination"],
        ["Chapter 6: Recommendations and Limitations of the Study", "To be updated after final pagination"],
        ["Bibliography / References", "To be updated after final pagination"],
        ["Appendix", "To be updated after final pagination"],
    ]
    add_table(doc, ["Section", "Page No."], toc_rows)

    doc.add_heading("LIST OF TABLES", level=1)
    add_table(
        doc,
            ["Table No.", "Title"],
        [
            ["Table 1", "Major Password Analysis Features"],
            ["Table 2", "Research Methodology Summary"],
            ["Table 3", "Project Modules and Responsibilities"],
            ["Table 4", "API Endpoint Summary"],
            ["Table 5", "Database and Authentication Design Decision"],
            ["Table 6", "Sample Password Analysis Results"],
            ["Table 7", "Testing Summary"],
        ],
    )

    doc.add_heading("LIST OF FIGURES", level=1)
    add_table(
        doc,
        ["Figure No.", "Title"],
        [
            ["Figure 1", "System Architecture of the Password Analyzer"],
            ["Figure 2", "Data Flow of Password Analysis"],
            ["Figure 3", "Use Case Diagram"],
            ["Figure 4", "Flowchart of Password Analysis"],
            ["Figure 5", "Web Interface of the Analyzer"],
            ["Figure 6", "Generated Password Output"],
            ["Figure 7", "Analysis Dashboard with Feedback"],
        ],
    )
    page_break(doc)


def abstract_paragraphs() -> list[str]:
    return [
        "Passwords continue to be one of the most widely used authentication mechanisms in web applications, "
        "mobile applications, online banking systems, learning portals, email services, and enterprise systems. "
        "Although modern authentication may include multi-factor verification and biometric checks, password "
        "security remains important because many users still create passwords that are short, predictable, reused, "
        "or based on personal information. A weak password can expose an account to dictionary attacks, brute-force "
        "attacks, password spraying, credential stuffing, and guessing attacks based on names, dates, phone numbers, "
        "or keyboard patterns.",
        "This project presents an AI-powered password strength analyzer developed as a practical BCA final-year "
        "project. The system evaluates a password by combining feature extraction, entropy estimation, logistic "
        "model prediction, pattern detection, explainable feedback, and a secure password generator. The analyzer "
        "does not merely count uppercase letters, lowercase letters, numbers, and symbols. It also identifies "
        "dictionary words, keyboard walks, repeated characters, date-like values, year patterns, numeric suffixes, "
        "and optional personal context such as name, email, and mobile number. This makes the feedback more useful "
        "than a simple rule-based meter.",
        "The project uses Python for backend logic and HTML, CSS, and JavaScript for the frontend dashboard. The "
        "backend exposes local API endpoints for analysis and secure password generation. The frontend uses AJAX "
        "requests, so analysis updates dynamically without page reload. The model uses engineered password features "
        "and a logistic probability function exposed through a predict_proba method. The final score combines model "
        "probability, entropy score, and pattern penalties. The interface displays AI model confidence, weak or "
        "strong prediction, entropy bits, effective entropy bits, estimated attack time, strengths, risks, "
        "improvement suggestions, and an AI explanation section.",
        "The system also includes a secure password generator implemented with Python's secrets module. The generated "
        "password contains lowercase characters, uppercase characters, digits, and symbols. The generated password is "
        "immediately analyzed by the same model so that the user can compare weak and strong examples. The application "
        "runs locally and does not store the password, name, email, or mobile number. This privacy-aware design is "
        "important because password analysis tools should avoid collecting sensitive input unnecessarily.",
        "The outcome of the project is a working model and dashboard that demonstrates practical knowledge of "
        "cybersecurity, web development, feature engineering, machine learning concepts, user interface design, and "
        "software testing. The project is suitable for academic demonstration because it is easy to run, transparent "
        "in its logic, and connected to real security practices such as entropy estimation, password strength meters, "
        "and secure random password generation.",
    ]


def extended_abstract_sections() -> list[tuple[str, list[str]]]:
    return [
        (
            "Background and Purpose",
            [
                "The present project is based on the continuing importance of password security in personal, academic, "
                "commercial, and institutional information systems. Passwords are often the first security boundary "
                "between an authorized user and a protected digital service. Even when organizations introduce "
                "multi-factor authentication or device-based verification, the password remains a key credential in "
                "many workflows. The weakness of this credential is that it is usually created by a human user. Human "
                "users naturally prefer memorable words, names, dates, and simple patterns. This makes many passwords "
                "easy to guess even when they appear to satisfy basic composition rules.",
                "The purpose of the AI-Powered Password Strength Analyzer is to develop a working model that evaluates "
                "passwords in a more informative and practical way. The system is not limited to checking whether a "
                "password contains a capital letter, a digit, or a symbol. It extracts a wider set of features and "
                "uses those features to estimate password strength. It also explains why the password is strong or "
                "weak. This is important because a security tool should not only reject weak passwords; it should help "
                "the user understand the reason for the weakness and suggest a safer alternative.",
                "The project was developed as a local web application for academic demonstration. The user enters a "
                "password in the browser, and the frontend sends the value to a local Python backend through an AJAX "
                "request. The backend performs feature extraction, applies the logistic model, calculates entropy, "
                "applies pattern penalties, and returns a structured JSON response. The interface then updates the "
                "score, color, risks, suggestions, model confidence, and AI explanation without reloading the page. "
                "This makes the project interactive and easy to demonstrate during viva.",
                "The system also includes a secure password generator. This generator uses Python's secrets module, "
                "which is designed for cryptographic randomness. The generated password includes lowercase letters, "
                "uppercase letters, digits, and symbols. After generation, the password is analyzed by the same model "
                "so the user can immediately see the difference between a human-created weak pattern and a randomly "
                "generated strong password. This supports both analysis and awareness.",
            ],
        ),
        (
            "Methodology",
            [
                "The project follows an applied software-development methodology. The problem was first identified as "
                "the limitation of simple password meters. The next step was to define the measurable features that "
                "could represent password strength. These features include length, uppercase count, lowercase count, "
                "digit count, special character count, number of unique characters, repeated-character ratio, keyboard "
                "patterns, simple sequences, dictionary words, year patterns, date-like values, suffix digits, and "
                "personal-information matches. The entropy value is calculated using the formula Entropy = L x log2(N), "
                "where L is the password length and N is the estimated character set size.",
                "After feature extraction, model-ready normalized values are prepared. These values are passed to a "
                "lightweight logistic model. The model is intentionally simple enough to be explained in an academic "
                "report and viva. It uses weights, a bias value, and a sigmoid function to produce a probability. The "
                "project exposes this through a predict_proba method that returns weak and strong class probabilities. "
                "The final score is not based only on the model; it also considers entropy and pattern penalties so "
                "that obvious human patterns do not receive unrealistic high scores.",
                "The training script creates synthetic weak, medium, and strong examples. Weak examples include common "
                "terms such as password, qwerty, admin, student, and name-based patterns with years or numeric suffixes. "
                "Medium examples include word-plus-number forms. Strong examples include longer random strings and "
                "multi-word passphrases. Synthetic training data is used because real leaked password datasets may "
                "contain sensitive information and are not suitable for a student project submission.",
                "The frontend methodology focuses on usability. The interface is divided into an input panel and an "
                "analysis panel. The input panel contains the password field, show/hide button, copy button, generator "
                "length control, generator button, and optional personal context fields. The analysis panel contains "
                "the animated score ring, AI model confidence, entropy details, attack-time estimate, strengths, risks, "
                "suggestions, and AI explanation. Red, orange, and green colors are used to communicate weak, medium, "
                "and strong scores visually.",
            ],
        ),
        (
            "Results and Interpretation",
            [
                "The working application successfully identifies common weak patterns. For example, a password such as "
                "password123 receives a very low score even though it includes letters and digits. The system detects "
                "the common word password, the numeric suffix, and the simple sequence. This result is important "
                "because it demonstrates that the analyzer is not fooled by superficial complexity. It recognizes that "
                "many attackers will try such patterns very early during dictionary or hybrid attacks.",
                "A password such as Sachin@2026 is also treated carefully when the optional name field contains Sachin. "
                "The analyzer detects the personal context and warns the user that the password contains personal "
                "information. This is a useful feature because attackers often know or guess names, email fragments, "
                "phone numbers, and meaningful years. The project therefore connects technical scoring with realistic "
                "human risk.",
                "Strong generated passwords receive high scores because they have sufficient length, multiple character "
                "classes, high entropy, no common dictionary hit, no keyboard walk, and no personal-information match. "
                "The secure generator gives the user a direct way to improve password quality. The copy button then "
                "allows the generated password to be copied to the clipboard, which improves usability during the demo.",
                "The AI explanation section improves transparency. A password meter that only displays a score may "
                "look mysterious to a user. This project explains the model prediction, confidence, entropy level, and "
                "pattern penalties. For weak passwords, it states that the model and pattern checks indicate weakness. "
                "For strong passwords, it states that entropy and character range are helping the password. This makes "
                "the system more educational and easier to defend during evaluation.",
            ],
        ),
        (
            "Security and Privacy",
            [
                "The project is designed with privacy as a central principle. It does not use a database because storing "
                "passwords or personal context values would create unnecessary risk. The password is analyzed in memory "
                "and returned as a score and explanation. The backend does not write the password to a file or table. "
                "The frontend also does not persist it in browser storage. This design decision is important enough to "
                "be documented because many student projects add a database without considering whether persistence is "
                "actually required.",
                "The optional name, email, and mobile fields are also not stored. They are used only as analysis context "
                "to detect whether the password contains personal information. For example, if the name field contains "
                "Sachin and the password contains the same name, the system reports a personal-context risk. This is "
                "not a login system and does not authenticate users. It is a stateless analysis tool. Therefore, there "
                "is no user registration, login, admin panel, or session management in the current implemented project.",
                "The secure password generator uses cryptographic randomness through Python's secrets module. This is "
                "more appropriate for password generation than ordinary pseudo-random methods intended for simulation "
                "or games. The generator also enforces a minimum length and includes all major character categories. "
                "The generated password is immediately passed to the analyzer so that the final dashboard output remains "
                "consistent with the same scoring system.",
                "The local API returns structured JSON and handles invalid JSON input with an error response. The server "
                "also sends no-store cache headers to reduce unnecessary browser caching of analysis responses. The "
                "application is designed for local educational use rather than public deployment. If deployed publicly, "
                "additional production controls such as HTTPS, request limits, stronger error handling, and logging "
                "redaction would be required.",
            ],
        ),
        (
            "Conclusion of Extended Abstract",
            [
                "The project demonstrates how cybersecurity awareness can be improved through a practical software "
                "model. It shows that a good password analyzer should combine mathematical estimation, machine-learning "
                "style probability, pattern recognition, and clear user feedback. It also shows that a project can be "
                "privacy-conscious by avoiding unnecessary storage. The final system is simple enough to run in a basic "
                "college lab environment and detailed enough to explain important security concepts.",
                "The expected academic value of the project lies in its integration of multiple BCA-relevant skills: "
                "Python programming, web development, data representation, algorithmic thinking, model training, unit "
                "testing, user-interface design, and cybersecurity reasoning. The report expansion should therefore "
                "focus not on adding unrelated features, but on explaining these implemented components deeply and "
                "professionally. This preserves the same project while meeting the word-count and formatting expectations "
                "of the university guidelines.",
            ],
        ),
    ]


def build_main_report(doc: Document) -> None:
    chapter_1(doc)
    chapter_2(doc)
    chapter_3(doc)
    chapter_4(doc)
    chapter_5(doc)
    chapter_6(doc)


def chapter_1(doc: Document) -> None:
    doc.add_heading("CHAPTER 1: INTRODUCTION TO THE TOPIC", level=1)
    for para in [
        "Authentication is the process of verifying that a user is who they claim to be. In most information "
        "systems, passwords are still the most common authentication secret. They are simple to implement, easy to "
        "understand, and widely supported across devices. However, their security depends heavily on user behavior. "
        "Many users choose passwords that are memorable but predictable, such as names, birth years, mobile numbers, "
        "sports terms, keyboard patterns, or simple dictionary words with digits added at the end.",
        "A password strength analyzer is a tool that helps users understand whether a password is likely to resist "
        "common guessing attacks. Traditional password meters often rely on basic rules: minimum length, at least one "
        "uppercase letter, at least one lowercase letter, at least one digit, and at least one special character. "
        "Such rules are easy to teach but can mislead users. For example, Password@123 satisfies many character "
        "composition rules but remains weak because it contains the common word password, a predictable capitalized "
        "shape, a symbol replacement, and a common numeric sequence.",
        "The project titled AI-Powered Password Strength Analyzer addresses this gap by combining a machine-learning "
        "style prediction model with explainable cybersecurity checks. The system is designed as a local web "
        "application. A user enters a password, optionally enters personal context fields, and receives a live "
        "strength score along with risks and suggestions. The interface also includes a secure password generator, a "
        "copy button, animated score ring, dynamic red-orange-green coloring, and a loading spinner that displays "
        "Analyzing password while the backend processes the request.",
        "The purpose of this project is not to store or manage real passwords. Instead, it is an educational security "
        "tool that demonstrates how passwords can be evaluated more intelligently. The analyzer calculates length, "
        "uppercase count, lowercase count, digit count, special character count, repeated characters, entropy, "
        "keyboard patterns, dictionary patterns, date patterns, and personal-information matches. These features are "
        "then used by a logistic model to estimate whether the password appears strong or weak.",
        "The entropy formula used in the project is Entropy = L x log2(N), where L is the password length and N is "
        "the estimated character set size. Entropy gives an approximate measure of the uncertainty faced by an "
        "attacker. However, entropy alone is not sufficient because human-created passwords often contain patterns. "
        "Therefore, the system combines entropy with pattern penalties and model confidence.",
        "The project is relevant to the field of computer applications because it involves software engineering, "
        "cybersecurity, web technologies, data analysis, and artificial intelligence concepts. It also helps create "
        "security awareness among users. A user who sees exactly why a password is weak is more likely to improve "
        "their behavior than a user who sees only a vague message.",
    ]:
        add_para(doc, para)

    doc.add_heading("1.1 Need of the Study", level=2)
    for para in [
        "The need for this study arises from the continued use of weak passwords in digital services. Even when "
        "applications enforce password complexity rules, users often follow predictable patterns to satisfy those "
        "rules with minimum effort. They may capitalize the first letter, add a symbol near the middle, and add a "
        "year or number at the end. Attackers are aware of these habits and include them in automated guessing "
        "strategies.",
        "A better password analyzer should do more than reward complexity symbols. It should detect human patterns, "
        "evaluate length and randomness, recognize personal information, and explain results in simple language. This "
        "project is built around that idea. It gives categorized output such as strengths, risks, suggestions, and AI "
        "explanation. This improves both usability and learning value.",
    ]:
        add_para(doc, para)

    doc.add_heading("1.2 Justification for Topic Selection", level=2)
    add_para(
        doc,
        "The topic was selected because password security is directly connected with daily digital life and the BCA "
        "curriculum. It allows the project to demonstrate Python programming, frontend design, backend APIs, feature "
        "engineering, local model prediction, and testing. It is also practical for a live viva because the examiner "
        "can type different passwords and immediately observe how the system responds.",
    )

    doc.add_heading("1.3 Background of Password-Based Authentication", level=2)
    for para in [
        "Password-based authentication is popular because it is inexpensive, familiar, and compatible with almost every "
        "digital platform. A password is a memorized secret known to the user and checked by the system. In theory, a "
        "password should be difficult for any other person or automated program to guess. In practice, users are often "
        "asked to remember many passwords across different services, so they choose patterns that are easy to recall. "
        "This behavior creates a gap between theoretical password strength and real password strength.",
        "The weakness of password authentication is not only a technical issue. It is also a human-computer interaction "
        "issue. When systems demand complexity without explanation, users may respond by creating predictable variants "
        "of familiar words. For example, a user may take a name, capitalize the first letter, replace a letter with a "
        "symbol, and add a year. Such a password may pass a basic policy, but it remains easy for attackers to guess "
        "because the transformation follows a common pattern.",
        "A strength analyzer therefore has an educational role. It should help users understand the difference between "
        "visible complexity and real unpredictability. Visible complexity means the password contains different character "
        "types. Real unpredictability means the password does not contain obvious words, personal details, or simple "
        "sequences. This project is designed around that distinction. It rewards useful length and variety while reducing "
        "the score for patterns that attackers commonly try.",
    ]:
        add_para(doc, para)

    doc.add_heading("1.4 Common Password Attack Methods", level=2)
    for para in [
        "A brute-force attack attempts many possible combinations until the correct password is found. The time required "
        "for brute force depends on password length, character set size, attacker speed, and the hashing method used by "
        "the target system. Long random passwords are difficult to brute force because the number of possible combinations "
        "grows rapidly. However, if a password is short or uses only digits, the search space becomes much smaller.",
        "A dictionary attack uses lists of common words, names, leaked passwords, and popular phrases. Attackers often "
        "combine dictionary terms with predictable modifications. They may add 123, 2024, 2025, or a symbol such as @. "
        "They may also test capitalized versions of words. This is why Password@123 is still weak even though it includes "
        "a capital letter, a symbol, and digits. The project includes a common-word check to demonstrate this risk.",
        "A keyboard-pattern attack tests strings that are easy to type because adjacent keys are used. Examples include "
        "qwerty, asdf, zxcv, 123456, and reversed sequences. These patterns are memorable for users and therefore common "
        "in real passwords. The feature engineering module detects keyboard walks by scanning known keyboard rows in both "
        "forward and reverse order. This is an important difference from a simple meter that only counts character types.",
        "A personal-information attack uses details connected to the user. Names, mobile numbers, email fragments, birth "
        "years, and important dates can often be guessed from social media, resumes, college portals, or public profiles. "
        "The optional name, email, and mobile fields in the project are included to demonstrate this risk. They are not "
        "stored and are not used for login. They simply allow the analyzer to check whether the password includes personal "
        "context supplied by the user for the current analysis.",
    ]:
        add_para(doc, para)

    doc.add_heading("1.5 Existing System and Its Limitations", level=2)
    for para in [
        "The existing approach in many websites is based on password rules. A user may be told that a password must have "
        "at least eight characters, one uppercase letter, one lowercase letter, one digit, and one symbol. Such policies "
        "are easy to implement, but they can produce a false sense of security. Users learn the minimum pattern required "
        "to satisfy the rule and then reuse similar passwords across different platforms.",
        "Another limitation of basic password meters is that they often provide poor feedback. A meter may show weak, "
        "medium, or strong without explaining the reason. If the user does not know why the password is weak, the user "
        "may make superficial changes. For example, adding an exclamation mark at the end may increase the appearance of "
        "complexity but does not remove the underlying dictionary word. A useful analyzer should identify the root cause.",
        "Many online password checkers also create privacy concerns. If a user types a real password into an unknown "
        "website, the password may be transmitted or logged. Even if the website claims to be safe, users cannot easily "
        "verify how the data is handled. This project avoids that risk by running locally and not storing passwords. The "
        "privacy design is an intentional part of the system, not an accidental omission.",
    ]:
        add_para(doc, para)

    doc.add_heading("1.6 Proposed System Overview", level=2)
    for para in [
        "The proposed system is a local AI-powered password strength analyzer. It combines three layers: feature "
        "engineering, model prediction, and explainable rule-based feedback. The feature layer extracts objective signals "
        "from the password. The model layer converts normalized features into a weak or strong probability. The feedback "
        "layer converts raw findings into human-readable strengths, risks, suggestions, and AI explanation points.",
        "The system has two main API operations. The /api/analyze endpoint accepts a password and optional context fields. "
        "It returns score, label, model probability, model confidence, entropy, effective entropy, attack-time estimates, "
        "risks, strengths, suggestions, explanation, and raw features. The /api/generate endpoint creates a secure random "
        "password and immediately analyzes it. These endpoints are used by the JavaScript frontend through AJAX.",
        "The frontend is designed as a dashboard rather than a plain form. The left panel contains user input and generator "
        "controls. The right panel contains analysis output. The animated score circle uses red, orange, and green colors "
        "for weak, medium, and strong results. A loading spinner displays Analyzing password during the request. The copy "
        "button allows the user to copy a generated password to the clipboard.",
    ]:
        add_para(doc, para)

    doc.add_heading("1.7 Scope of the Project", level=2)
    for para in [
        "The scope of the project is limited to password analysis and secure password generation. It does not implement "
        "account registration, login, role-based access control, admin dashboards, or a database. These features are not "
        "required for the current objective because the project is not a user-management system. It is a focused security "
        "analysis tool. Adding login or a database would not improve the core password-strength model and could introduce "
        "unnecessary privacy risk.",
        "The system is intended for academic demonstration, security awareness, and local experimentation. It can help a "
        "student or evaluator compare weak and strong passwords, observe how entropy changes, and understand why common "
        "patterns reduce strength. It should not be treated as a complete enterprise password-auditing tool because it "
        "does not connect to a real leaked-password database and does not test against every possible attack strategy.",
    ]:
        add_para(doc, para)

    doc.add_heading("1.8 Significance of the Study", level=2)
    for para in [
        "The project is significant because it connects academic programming with a real security problem. Many students "
        "learn Python, HTML, CSS, JavaScript, and basic machine-learning concepts separately. This project combines them "
        "into one usable application. The backend handles analysis, the model performs prediction, the frontend displays "
        "results interactively, and the tests verify important cases. This makes it suitable for a final-year project.",
        "The study also shows that cybersecurity education is improved when tools explain decisions. A user who sees a "
        "red score and a message such as Contains common words: password can understand the specific weakness. A user who "
        "sees Replace dates and years with random words or characters receives an actionable improvement path. This is "
        "more useful than a generic message such as Please choose a stronger password.",
    ]:
        add_para(doc, para)

    add_table(
        doc,
        ["Feature", "Purpose in Analysis"],
        [
            ["Length", "Longer passwords increase the search space for brute-force attacks."],
            ["Uppercase count", "Measures use of capital letters and helps identify character variety."],
            ["Lowercase count", "Measures normal alphabetic content and helps calculate character classes."],
            ["Digit count", "Detects numeric usage and numeric suffix patterns."],
            ["Special character count", "Detects symbols that increase possible character set size."],
            ["Repeated characters", "Detects weak patterns such as aaa, 111, or repeated blocks."],
            ["Entropy", "Estimates uncertainty using Entropy = L x log2(N)."],
            ["Personal context", "Detects names, email parts, and mobile numbers inside the password."],
        ],
        "Table 1: Major Password Analysis Features",
    )

    add_caption(doc, "Figure 1: System Architecture of the Password Analyzer")
    add_architecture_table(doc)

    add_caption(doc, "Figure 2: Data Flow of Password Analysis")
    add_data_flow_table(doc)
    page_break(doc)


def add_architecture_table(doc: Document) -> None:
    add_table(
        doc,
        ["Layer", "Component", "Description"],
        [
            ["Presentation", "HTML, CSS, JavaScript", "Provides input fields, generator button, score dashboard, and feedback cards."],
            ["Communication", "AJAX Fetch API", "Sends JSON requests to local backend without page reload."],
            ["Backend", "Python HTTP API", "Receives password input, calls analyzer, and returns JSON response."],
            ["AI Logic", "Feature extraction and logistic model", "Converts password into numeric features and predicts weak or strong class."],
            ["Security Logic", "Pattern checks and generator", "Detects common risks and creates secure random passwords with secrets."],
        ],
    )


def add_data_flow_table(doc: Document) -> None:
    add_table(
        doc,
        ["Step", "Process"],
        [
            ["1", "User enters or generates a password in the browser."],
            ["2", "JavaScript sends password and optional context to /api/analyze."],
            ["3", "Feature engineering module extracts length, counts, entropy, and patterns."],
            ["4", "Model module returns weak and strong probabilities through predict_proba."],
            ["5", "Analyzer combines model probability, entropy, and penalties into final score."],
            ["6", "Frontend displays score, colors, suggestions, attack estimate, and AI explanation."],
        ],
    )


def chapter_2(doc: Document) -> None:
    doc.add_heading("CHAPTER 2: REVIEW OF LITERATURE", level=1)
    for para in [
        "Password strength evaluation has developed from simple composition rules toward more realistic estimation "
        "methods. Earlier systems commonly evaluated a password by checking the presence of lowercase letters, "
        "uppercase letters, digits, and symbols. These rules are known as LUDS rules. They are easy to implement but "
        "do not fully capture how attackers guess passwords. Research and security guidance have shown that users "
        "often respond to complexity rules with predictable transformations rather than truly random secrets.",
        "Modern security guidance emphasizes password length, avoidance of known compromised passwords, user-friendly "
        "strength meters, and protection against online attack methods such as brute force and credential stuffing. "
        "The OWASP Authentication Cheat Sheet recommends proper password strength controls and password strength "
        "meters as part of secure authentication design. It also discusses automated attacks such as brute force, "
        "credential stuffing, and password spraying.",
        "NIST Digital Identity Guidelines describe passwords as memorized secrets and define entropy as uncertainty "
        "in bits. A password strength analyzer can use entropy as a helpful mathematical signal, but entropy must be "
        "interpreted carefully. A random 16-character password and a human-made 16-character password may have the "
        "same length but very different practical resistance if the human password contains common words, names, or "
        "dates.",
        "The zxcvbn password strength estimator, presented by Wheeler at USENIX Security 2016, is important in this "
        "area because it highlighted the weakness of relying only on character-class rules. The paper argues that "
        "realistic password strength estimation should consider patterns and guessing attacks. This project is not a "
        "copy of zxcvbn, but it is inspired by the general idea that password feedback should be pattern-aware and "
        "user-friendly.",
        "Password storage literature also influences the project indirectly. Although this analyzer does not store "
        "passwords, password security must be understood within the larger authentication environment. The OWASP "
        "Password Storage Cheat Sheet explains that passwords should not be stored in plaintext and should be "
        "protected with slow hashing algorithms such as Argon2id, scrypt, bcrypt, or PBKDF2. This reinforces the "
        "principle that password handling should be privacy-conscious.",
        "Machine learning and feature engineering can improve a password analyzer by allowing multiple features to "
        "contribute to a prediction. Instead of using a single rule, a model can combine length, entropy, character "
        "variety, repetition, dictionary indicators, keyboard indicators, and personal-context matches. In this "
        "project, a lightweight logistic model is used because it is understandable, easy to run locally, and suitable "
        "for academic demonstration without external dependencies.",
    ]:
        add_para(doc, para)

    doc.add_heading("2.1 Review Summary", level=2)
    add_bullets(
        doc,
        [
            "Simple composition rules are not enough for realistic password evaluation.",
            "Length and entropy are useful, but they must be combined with pattern detection.",
            "Password meters should give understandable feedback rather than only accepting or rejecting a password.",
            "Privacy is important; password analysis should avoid unnecessary storage or external transmission.",
            "A lightweight AI-style model can be practical for educational systems when paired with explainable rules.",
        ],
    )

    doc.add_heading("2.2 Password Composition Rules", level=2)
    for para in [
        "Password composition rules became common because they are easy for systems to check. A program can quickly "
        "verify whether a password contains uppercase letters, lowercase letters, numbers, and symbols. These rules were "
        "introduced to discourage very simple passwords, but they do not measure unpredictability directly. Users often "
        "learn predictable ways to satisfy them, such as adding a capital letter at the beginning and a number at the end.",
        "For an academic password analyzer, composition rules are still useful as features, but they should not be treated "
        "as the complete answer. This project uses character classes as part of the feature set. The system counts lower "
        "case letters, uppercase letters, digits, and special characters. It also calculates how many different character "
        "classes are present. However, the final score depends on additional signals such as entropy, dictionary hits, "
        "keyboard walks, repetition, personal context, and model probability.",
        "This distinction is important in the literature review because it explains why the project is not just a simple "
        "password checker. It is a feature-based analyzer. Character classes are treated as one category of evidence "
        "among many, rather than as a guarantee of strength. This approach is more realistic and easier to explain during "
        "evaluation because it shows awareness of the limitations of older password policies.",
    ]:
        add_para(doc, para)

    doc.add_heading("2.3 Entropy-Based Evaluation", level=2)
    for para in [
        "Entropy is a mathematical way to estimate uncertainty. In password analysis, entropy is often expressed in bits. "
        "A higher entropy value means that an attacker has more possible combinations to search. The simple formula used "
        "in this project is Entropy = L x log2(N), where L is the password length and N is the estimated character set "
        "size. A password using only digits has a smaller N than a password using lowercase, uppercase, digits, and symbols.",
        "Entropy is valuable because it rewards length and variety. A longer password increases the search space, and a "
        "larger character set increases the number of choices per position. However, entropy calculated only from length "
        "and character classes can overestimate human-created passwords. Password@123 has high theoretical character-set "
        "coverage, but it contains the word password and a common numeric sequence. Therefore, the project uses entropy as "
        "one signal and then applies pattern penalties.",
        "The report should explain entropy carefully because it is a central technical concept in the project. Entropy is "
        "not presented as a perfect measure. It is presented as a useful approximation that becomes stronger when combined "
        "with pattern-aware analysis. This balanced explanation is more accurate than claiming that the formula alone can "
        "determine real-world strength.",
    ]:
        add_para(doc, para)

    doc.add_heading("2.4 Pattern-Aware Password Meters", level=2)
    for para in [
        "Pattern-aware password meters attempt to evaluate how people actually create passwords. Instead of assuming that "
        "every character is chosen independently and randomly, they look for words, names, keyboard paths, dates, repeated "
        "characters, and common substitutions. This approach is closer to attacker behavior because attackers do not only "
        "try random combinations. They try likely human patterns first.",
        "The project follows this idea by checking a custom list of common words and names, detecting leet substitutions "
        "such as @ for a or 0 for o, scanning keyboard rows, and identifying year and date patterns. These checks make the "
        "system more sensitive to weak but visually complex passwords. A user can therefore receive feedback such as Remove "
        "keyboard patterns or Replace common word(s) password with unrelated words.",
        "The zxcvbn password strength estimator is a well-known example of pattern-aware thinking. This project does not "
        "attempt to reimplement that tool. Instead, it applies the same broad lesson at a level suitable for a BCA project: "
        "password strength is affected by recognizable patterns, and a good analyzer should explain those patterns to the "
        "user.",
    ]:
        add_para(doc, para)

    doc.add_heading("2.5 Machine Learning and Feature Engineering", level=2)
    for para in [
        "Machine learning becomes useful when multiple measurable features must be combined into a prediction. In this "
        "project, the model is intentionally lightweight. A logistic model is used because it is explainable and can run "
        "without third-party packages. The model takes normalized feature values and produces weak and strong class "
        "probabilities. This gives the project an AI-oriented component while keeping the implementation transparent.",
        "Feature engineering is the process of converting raw password text into numeric signals. The raw password itself "
        "is not used as a stored training record. Instead, the system extracts properties such as length, entropy, number "
        "of character classes, uniqueness ratio, repetition, sequence count, keyboard pattern count, dictionary-word hits, "
        "and personal-context hits. These values are safer and more explainable than storing real passwords.",
        "The model is trained with synthetic examples. This is a practical design choice for a student project because real "
        "password datasets may be unethical or unsafe to use. Synthetic weak examples represent common risky patterns, "
        "medium examples represent partially improved patterns, and strong examples represent long random strings or "
        "passphrases. The training script then applies gradient descent to learn weights for the selected features.",
    ]:
        add_para(doc, para)

    doc.add_heading("2.6 Privacy in Password Analysis Tools", level=2)
    for para in [
        "Privacy is an important theme in password-strength tools. A user should be careful about typing a real password "
        "into an unknown online service. If the service logs the password, sends it to a third party, or stores it in a "
        "database, the tool creates the same risk it claims to reduce. For this reason, the project is designed as a local "
        "application and avoids persistence.",
        "The absence of a database is therefore a deliberate security decision. A database would be useful for a login "
        "portal, admin panel, or analytics system, but this project is not an account-management system. Its objective is "
        "to analyze a password and immediately return feedback. No account is created, no session is stored, and no user "
        "record is saved. This keeps the system focused and reduces sensitive data exposure.",
        "The report should explicitly explain this point because some project formats ask for database tables by default. "
        "In this project, database connectivity is not missing due to negligence; it is not required by the functional "
        "scope. The correct documentation is a privacy-first design decision and a future-scope note for optional analytics "
        "or administrative features.",
    ]:
        add_para(doc, para)

    doc.add_heading("2.7 Secure Random Password Generation", level=2)
    for para in [
        "A password analyzer can identify weaknesses, but users also need a way to create stronger passwords. The secure "
        "password generator solves this by producing a random password with sufficient length and character variety. The "
        "generator uses Python's secrets module, which is intended for security-sensitive randomness. This is more "
        "appropriate than using ordinary random functions that are designed for simulations.",
        "The generator includes at least one lowercase letter, one uppercase letter, one digit, and one symbol. It then "
        "fills the remaining positions from the combined alphabet and securely shuffles the result. The length is clamped "
        "between a safe minimum and maximum so that an invalid input cannot create a very short password. The generated "
        "password is analyzed immediately, demonstrating that the generator and analyzer are connected parts of the same "
        "workflow.",
    ]:
        add_para(doc, para)

    doc.add_heading("2.8 Literature Review Conclusion", level=2)
    for para in [
        "The literature and technical guidance support the main direction of the project. Password strength should not be "
        "judged only by simple composition rules. Entropy is useful, but it must be interpreted with awareness of human "
        "patterns. Machine-learning style scoring can combine multiple features, but the result should remain explainable "
        "to the user. Privacy is also important because passwords are sensitive secrets.",
        "Based on this review, the project adopts a hybrid design. It combines entropy, feature engineering, logistic "
        "prediction, pattern detection, secure generation, and human-readable feedback. This design is appropriate for a "
        "BCA final-year project because it demonstrates practical implementation while remaining technically transparent.",
    ]:
        add_para(doc, para)


def chapter_3(doc: Document) -> None:
    doc.add_heading("CHAPTER 3: RESEARCH OBJECTIVES AND METHODOLOGY", level=1)
    doc.add_heading("3.1 Research Objectives", level=2)
    add_bullets(
        doc,
        [
            "To design and develop a working password strength analyzer using Python and web technologies.",
            "To extract meaningful password features such as length, character counts, repeated characters, entropy, and weak patterns.",
            "To apply an AI-style logistic prediction model for weak or strong password classification.",
            "To provide intelligent improvement suggestions for weak passwords.",
            "To generate secure passwords using cryptographic randomness.",
            "To create a user-friendly dashboard that updates dynamically without page reload.",
        ],
    )

    doc.add_heading("3.2 Research Problem", level=2)
    add_para(
        doc,
        "The research problem is that many password meters provide either overly simple feedback or misleading scores. "
        "A password may pass composition rules but still be predictable. The project investigates how a password can "
        "be evaluated through combined signals: entropy, engineered features, model probability, and explainable "
        "pattern detection.",
    )

    doc.add_heading("3.3 Hypotheses", level=2)
    add_bullets(
        doc,
        [
            "Null hypothesis: A feature-based AI password analyzer does not provide more useful feedback than a simple rule-based meter.",
            "Alternative hypothesis: A feature-based AI password analyzer provides more useful feedback by combining entropy, model probability, and pattern-specific suggestions.",
        ],
    )

    doc.add_heading("3.4 Research Design", level=2)
    add_para(
        doc,
        "The study follows an applied and descriptive research design. It is applied because the output is a working "
        "software model. It is descriptive because the system describes password strength through measurable features, "
        "scores, risks, and suggestions. The project does not require a large external survey because its primary "
        "focus is software development and model demonstration.",
    )

    add_table(
        doc,
        ["Methodology Item", "Project Decision"],
        [
            ["Research Design", "Applied and descriptive software project"],
            ["Type of Data Used", "Synthetic password examples and user-entered test cases"],
            ["Data Collection Method", "Generated weak, medium, and strong password patterns"],
            ["Data Collection Instrument", "Python training script and browser input form"],
            ["Sample Size", "Synthetic training examples plus live test passwords"],
            ["Sampling Technique", "Non-probability purposive sampling of password patterns"],
            ["Data Analysis Tool", "Python feature extraction, entropy calculation, and logistic model"],
            ["Frontend Method", "HTML, CSS, JavaScript, and AJAX Fetch API"],
        ],
        "Table 2: Research Methodology Summary",
    )

    doc.add_heading("3.5 System Modules", level=2)
    add_table(
        doc,
        ["Module", "Responsibility"],
        [
            ["Feature Engineering", "Extracts numeric password features including length, character classes, entropy, repeated characters, dictionary hits, and personal context."],
            ["Model", "Loads logistic weights and returns weak/strong probabilities using predict_proba."],
            ["Analyzer", "Combines model score, entropy score, and penalties into final score and label."],
            ["Suggestions", "Creates targeted improvement suggestions for weak passwords."],
            ["Explainability", "Explains the model decision, confidence, entropy, and pattern penalties."],
            ["Generator", "Creates secure random passwords using Python secrets."],
            ["Frontend", "Displays password input, generator, score ring, dynamic colors, spinner, metrics, and feedback."],
        ],
        "Table 3: Project Modules and Responsibilities",
    )

    doc.add_heading("3.6 Tools and Technologies", level=2)
    add_bullets(
        doc,
        [
            "Python 3 for backend logic and model execution.",
            "Python standard library HTTP server for local API endpoints.",
            "HTML, CSS, and JavaScript for browser interface.",
            "AJAX Fetch API for dynamic password analysis without page reload.",
            "Python unittest for test verification.",
            "JSON file for storing trained logistic model weights.",
        ],
    )

    doc.add_heading("3.7 Entropy Formula", level=2)
    add_para(doc, "Entropy is calculated using the formula:")
    p = doc.add_paragraph(style="Code Block")
    p.add_run("Entropy = L x log2(N)")
    add_para(
        doc,
        "In this formula, L represents the password length and N represents the estimated character set size. For "
        "example, a password using lowercase letters, uppercase letters, digits, and symbols has a larger N than a "
        "password using only digits. Higher entropy generally means a larger search space, but pattern checks are "
        "still required because human-created passwords are rarely fully random.",
    )

    doc.add_heading("3.8 API Design and Backend Workflow", level=2)
    for para in [
        "The backend is implemented in app.py using Python's standard library HTTP server. The server serves static files "
        "from the static folder and handles JSON POST requests. Two API endpoints are implemented: /api/analyze and "
        "/api/generate. The backend is intentionally small and local because the main objective of the project is the "
        "working model, not production deployment. This also keeps the project easy to run in a college lab environment.",
        "When the user enters a password, JavaScript sends a JSON payload containing the password and optional personal "
        "context fields. The backend reads the request body, parses JSON, and passes the password to the analyzer. If the "
        "request is invalid JSON, the server returns a JSON error response. If the route is not recognized, the server "
        "returns a not found error. This demonstrates basic API validation and structured error handling.",
        "The /api/generate endpoint accepts a desired length. It calls generate_secure_password, receives the generated "
        "password, and then analyzes the generated value using the same analyzer pipeline. The response contains the "
        "password, its length, and the full analysis object. This design avoids duplicating frontend logic because the "
        "same render function can display results from typed passwords and generated passwords.",
    ]:
        add_para(doc, para)

    add_table(
        doc,
        ["Endpoint", "Method", "Input", "Output", "Purpose"],
        [
            ["/api/analyze", "POST", "password and optional context", "score, prediction, entropy, risks, suggestions, explanation", "Analyze a typed password."],
            ["/api/generate", "POST", "desired password length", "generated password and analysis object", "Create and analyze a secure password."],
        ],
        "Table 4: API Endpoint Summary",
    )

    doc.add_heading("3.9 Database and Authentication Design Decision", level=2)
    for para in [
        "The project does not include database connectivity. This is a deliberate design decision rather than a missing "
        "implementation. The application analyzes passwords and returns feedback. It does not need to create user accounts, "
        "store password history, maintain profiles, or save reports. Storing passwords or personal context would increase "
        "risk and would conflict with the privacy principle of the project.",
        "There is also no login or registration flow. The system is not an authentication portal; it is a password-strength "
        "analysis tool. A login system would require password storage, hashing, session handling, and possibly user roles. "
        "Those features belong to a different type of project. In this project, the correct workflow is stateless analysis: "
        "the user enters data, the server calculates feedback, and the response is displayed immediately.",
        "There is no admin panel in the implemented system. An admin panel could be part of a future version if the project "
        "were expanded into a classroom awareness platform that stores anonymous test statistics. For the current version, "
        "the absence of an admin panel keeps the project focused on the password analyzer model and reduces complexity.",
    ]:
        add_para(doc, para)

    add_table(
        doc,
        ["Requirement Area", "Current Project Status", "Reason"],
        [
            ["Database tables", "Not implemented", "No password or personal data is stored."],
            ["Login/Register", "Not implemented", "The app is an analyzer, not an account system."],
            ["Authentication sessions", "Not implemented", "No user identity is required for local analysis."],
            ["Admin panel", "Not implemented", "No stored records or users need administration."],
            ["Future analytics database", "Possible future scope", "Could store anonymous aggregate results only."],
        ],
        "Table 5: Database and Authentication Design Decision",
    )

    doc.add_heading("3.10 Use Case Diagram Description", level=2)
    add_caption(doc, "Figure 3: Use Case Diagram")
    add_table(
        doc,
        ["Actor", "Use Case", "Description"],
        [
            ["User", "Enter password", "Types a password into the analyzer input field."],
            ["User", "Show or hide password", "Temporarily reveals or masks the password field."],
            ["User", "Generate password", "Requests a secure password of selected length."],
            ["User", "Copy password", "Copies generated or typed password to clipboard."],
            ["User", "Enter optional context", "Provides name, email, or mobile for personal-information detection."],
            ["System", "Analyze password", "Extracts features, predicts strength, calculates score, and returns feedback."],
            ["System", "Explain result", "Shows model confidence, entropy interpretation, and pattern penalties."],
        ],
    )
    add_para(
        doc,
        "The use case diagram for this project has one primary human actor: the user. The system actor is the local "
        "password analyzer. The user can enter a password, generate a password, copy the password, show or hide it, and "
        "optionally provide context values. The system analyzes the password and displays results. There is no admin actor "
        "because the implemented project does not include stored records or administrative management.",
    )

    doc.add_heading("3.11 Data Flow Diagram Description", level=2)
    add_para(
        doc,
        "The Level 0 DFD can be described as a single process named Password Strength Analysis System. The user provides "
        "password input and optional context. The system returns score, label, model confidence, entropy, risk messages, "
        "suggestions, and AI explanation. There is no external database data store. The only stored model artifact is the "
        "local JSON file containing model weights.",
    )
    add_para(
        doc,
        "The Level 1 DFD separates the process into frontend input handling, backend request handling, feature extraction, "
        "model prediction, scoring and explanation, and frontend rendering. Password data flows through memory only. The "
        "frontend sends it to the local backend, the backend transforms it into features, and the response is returned as "
        "analysis output. This design should be represented in the report as a stateless data flow.",
    )

    doc.add_heading("3.12 Flowchart of Password Analysis", level=2)
    add_caption(doc, "Figure 4: Flowchart of Password Analysis")
    add_table(
        doc,
        ["Step", "Flowchart Action", "Decision / Output"],
        [
            ["1", "Start application", "Server runs on localhost."],
            ["2", "User enters or generates password", "Input is available in browser."],
            ["3", "AJAX request sent", "Request goes to /api/analyze or /api/generate."],
            ["4", "Validate request", "Invalid JSON returns error; valid JSON continues."],
            ["5", "Extract features", "Length, counts, entropy, and patterns are calculated."],
            ["6", "Predict probability", "Model returns weak and strong probabilities."],
            ["7", "Apply scoring logic", "Entropy, model score, and penalties are combined."],
            ["8", "Generate feedback", "Risks, strengths, suggestions, and explanation are prepared."],
            ["9", "Render result", "Frontend updates dashboard without page reload."],
            ["10", "End", "User can test another password."],
        ],
    )

    doc.add_heading("3.13 Validation Logic", level=2)
    for para in [
        "Validation occurs in multiple layers. The frontend controls the generator length input with a minimum of 12 and "
        "a maximum of 32. The backend also clamps the generated password length between the same safe range. This prevents "
        "very short generated passwords even if a user changes the frontend value manually.",
        "The analyzer accepts an empty password but returns a safe empty result with score zero and a message asking the "
        "user to create a password or passphrase. This is better than failing with an exception. The backend also catches "
        "invalid JSON and returns an error object. These validation choices make the demo more stable.",
        "The feature extraction layer normalizes text for common-word and personal-context detection. It removes non-alphanumeric "
        "characters and applies simple leet translation. This allows patterns such as p@ssword or pa55word to be treated "
        "closer to password. The validation and normalization logic supports more realistic analysis.",
    ]:
        add_para(doc, para)

    doc.add_heading("3.14 Main Algorithms", level=2)
    for para in [
        "The first algorithm is feature extraction. It scans the password and calculates direct counts such as length, "
        "lowercase characters, uppercase characters, digits, symbols, and unique characters. It then calculates derived "
        "features such as entropy, repeated-character ratio, longest repeated run, sequence triplets, keyboard walks, "
        "suffix digits, common-word hits, and context matches.",
        "The second algorithm is logistic prediction. The normalized feature vector is multiplied by learned weights and "
        "combined with a bias value. The result is passed through a sigmoid function, producing a probability between 0 "
        "and 1. The project exposes this result through predict_proba so the report can describe the model in familiar "
        "machine-learning terminology.",
        "The third algorithm is final scoring. The model probability is converted into a model score, entropy is converted "
        "into an entropy score, and pattern penalties are subtracted. The final score is bounded between 0 and 100. This "
        "prevents unusual inputs from producing impossible values. The score is then mapped to labels such as Very Weak, "
        "Weak, Fair, Strong, and Excellent.",
        "The fourth algorithm is suggestion generation. The suggestion module checks the same raw features and creates "
        "specific messages. If the password is short, it recommends increasing length. If character classes are missing, "
        "it recommends adding the missing classes. If common words, keyboard walks, dates, or personal details are found, "
        "it recommends removing them. This makes the feedback intelligent and targeted.",
    ]:
        add_para(doc, para)

    doc.add_heading("3.15 ER Diagram Applicability", level=2)
    add_para(
        doc,
        "An Entity Relationship Diagram is normally used when a project has persistent entities such as User, Admin, "
        "PasswordRecord, Report, or LoginSession. The current project intentionally does not persist such entities. "
        "Therefore, a conventional ER diagram is not applicable to the implemented version. For documentation purposes, "
        "the report can state that the data model is stateless and request-based. The temporary data objects are Password "
        "Input, Optional Context, Feature Bundle, Model Result, and Analysis Response. These exist only during processing "
        "and are not saved to a database.",
    )
    page_break(doc)


def chapter_4(doc: Document) -> None:
    doc.add_heading("CHAPTER 4: DATA ANALYSIS, RESULTS, AND INTERPRETATION", level=1)
    analyzer = PasswordStrengthAnalyzer.from_default_model()
    generated = generate_secure_password(18)

    doc.add_heading("4.1 Implementation Overview", level=2)
    for para in [
        "This chapter explains how the implemented project works internally and how the results produced by the system "
        "should be interpreted. The project is a working local web application. It has a Python backend, a browser-based "
        "frontend, a trained logistic model stored as JSON, feature extraction functions, password-improvement logic, "
        "and a secure generator. These parts are connected into one workflow so that a user can type or generate a "
        "password and immediately receive a strength score, AI prediction, entropy value, estimated attack time, risks, "
        "suggestions, and explanation.",
        "The implementation is intentionally lightweight because the project objective is to demonstrate a working "
        "password-strength model, not to build a large enterprise authentication platform. The backend uses Python's "
        "standard library HTTP server rather than a heavy framework. This makes the project simple to run on a college "
        "computer. At the same time, the internal design remains modular. Feature extraction, model prediction, analysis, "
        "suggestions, explainability, and generation are placed in separate Python files under the password_ai package.",
        "The system should be understood as an educational security tool. It does not accept a login, create user "
        "accounts, store passwords, or save personal details. This matters because passwords are sensitive values. A "
        "password analyzer that stores tested passwords would create an unnecessary risk. The implemented design avoids "
        "that risk by processing values in memory and returning the result to the browser immediately.",
        "The frontend communicates with the backend using AJAX through the JavaScript Fetch API. This means the page does "
        "not reload every time analysis is performed. The dashboard changes dynamically: the score ring color changes "
        "based on strength, the spinner appears while the request is processing, the entropy and crack-time sections are "
        "updated, and suggestion cards are replaced with the latest feedback. This improves the demonstration value of "
        "the project because users can test many examples quickly.",
    ]:
        add_para(doc, para)

    doc.add_heading("4.2 Module-Wise Implementation", level=2)
    for para in [
        "The app.py file is the entry point of the application. It creates the local HTTP server, serves the static files, "
        "and handles the API routes. When the browser requests the main page, the server returns static/index.html. When "
        "the browser sends an analysis request, app.py reads the JSON payload, calls the password analyzer, and returns a "
        "JSON response. The file also handles the generator route, error responses, and content types for static files.",
        "The password_ai/feature_engineering.py file contains the feature extraction functions. This is one of the most "
        "important files in the project because the AI model depends on these features. The file calculates direct "
        "measurements such as length, uppercase count, lowercase count, digit count, symbol count, and unique character "
        "count. It also calculates derived indicators such as entropy, repeated-character ratio, keyboard walk count, "
        "common-word hits, suffix digits, date-like patterns, and context matches.",
        "The password_ai/model.py file contains the lightweight logistic prediction model. The model loads weights and "
        "feature names from data/model.json. The predict_probability function applies the logistic formula to a feature "
        "vector and returns the probability that the password belongs to the strong class. The predict_proba function "
        "returns both weak and strong probabilities so the frontend can show AI confidence instead of a fixed or fake "
        "percentage.",
        "The password_ai/analyzer.py file is the central decision layer. It calls feature extraction, calls the model, "
        "calculates a final score, assigns a label, estimates attack time, builds strengths and risks, requests "
        "suggestions, and prepares the explanation. This file connects the technical model with the user-facing result. "
        "It also ensures that results are bounded and readable. For example, the final score is kept between 0 and 100.",
        "The password_ai/suggestions.py file creates intelligent improvement suggestions. The function does not return "
        "only a generic instruction such as use a stronger password. Instead, it checks the detected weakness and gives "
        "specific guidance. If the password is short, it recommends increasing length. If common words are detected, it "
        "recommends avoiding dictionary words. If repeated characters are found, it recommends replacing repetition with "
        "a longer phrase. If personal information is found, it recommends removing names, email parts, phone numbers, or "
        "years.",
        "The password_ai/explainability.py file prepares a plain-language explanation of the analysis result. This is "
        "important because AI results can appear mysterious if the project only displays a number. The explanation links "
        "the prediction to visible evidence such as entropy, model confidence, detected patterns, and applied penalties. "
        "This makes the system more suitable for academic viva because the student can explain why the score changed.",
        "The password_ai/generator.py file creates strong passwords using Python's secrets module. It guarantees character "
        "variety by including lowercase, uppercase, digit, and symbol characters. It also clamps the password length to "
        "a safe range. After generation, the password is shuffled securely and sent back to the analyzer so that the "
        "same dashboard can explain the generated password.",
        "The static/index.html, static/styles.css, and static/app.js files form the frontend. The HTML file defines the "
        "input layout and dashboard sections. The CSS file controls the professional two-panel visual design, score ring, "
        "cards, colors, spacing, and responsive behavior. The JavaScript file controls events such as typing, show or "
        "hide password, copying to clipboard, sending AJAX requests, showing the loading spinner, and rendering results.",
    ]:
        add_para(doc, para)

    doc.add_heading("4.3 Feature Extraction Function Analysis", level=2)
    for para in [
        "Feature extraction is the process of converting the raw password string into numeric and logical values that can "
        "be used by the model and scoring engine. A password is text, but a machine-learning model cannot directly "
        "understand text in the same way a human reads it. Therefore, the system measures useful properties and converts "
        "them into a structured feature bundle. This project extracts both composition features and pattern features.",
        "The length feature is the simplest but most important feature. Longer passwords usually create a larger search "
        "space because an attacker must try more possible combinations. However, length alone is not enough. A long value "
        "such as passwordpasswordpassword may still be weak because it is based on repetition and dictionary words. The "
        "analyzer therefore treats length as one signal, not as the complete answer.",
        "Uppercase count, lowercase count, digit count, and special character count measure character variety. These "
        "values help the system estimate the possible character set size used by the password. A password using only "
        "lowercase letters has a smaller estimated character set than a password using lowercase, uppercase, digits, and "
        "symbols. This directly affects the entropy formula used in the project.",
        "The repeated-character features detect weak construction such as aaa, 1111, or !!!!!. Repetition is common in "
        "human-created passwords because it is easy to remember and easy to type. The project calculates the longest "
        "repeated run and a repeated-character ratio. These values help the analyzer warn the user when the password is "
        "not as random as it appears.",
        "The entropy calculation uses the formula Entropy = L x log2(N), where L is the password length and N is the "
        "estimated character set size. The feature extraction function estimates N based on which character classes are "
        "present. If lowercase letters are present, their alphabet size contributes to N. If uppercase letters, digits, "
        "or symbols are present, they also increase N. The final entropy value is expressed in bits.",
        "Entropy is useful because it represents the theoretical search space, but it is not perfect for human-created "
        "passwords. The formula assumes that characters are chosen randomly from the available set. Many users do not "
        "choose characters randomly. They use names, words, years, keyboard walks, or simple substitutions. For that "
        "reason, the analyzer combines entropy with pattern detection and model prediction instead of using entropy alone.",
        "Common-word detection checks whether the password contains weak terms such as password, admin, welcome, qwerty, "
        "or other predictable fragments. The implementation also applies simple normalization and leet translation, so "
        "some substitutions such as @ for a or 5 for s can still be detected. This is important because users often "
        "believe that p@ssword is strong, even though attackers commonly test such variations.",
        "Keyboard-walk detection identifies sequences such as qwerty, asdf, or 12345. These patterns are especially weak "
        "because attackers include them in dictionaries and rule-based guessing tools. The analyzer counts such patterns "
        "and includes them in the risk output. A password containing a keyboard walk may receive a lower score even if "
        "it includes multiple character classes.",
        "Personal-context matching is another major feature. The frontend may collect optional name, email, and mobile "
        "values. These values are not required, and they are not stored. When provided, the backend checks whether parts "
        "of that information appear inside the password. This makes the analyzer more realistic because attackers often "
        "try passwords based on a person's name, date, phone number, or email username.",
        "The feature extraction module returns both raw features and normalized features. Raw features are useful for "
        "human-readable suggestions and explanations. Normalized features are useful for the model because machine-learning "
        "weights work better when input values are scaled into comparable ranges. This separation makes the code clearer "
        "and easier to explain in the project report.",
    ]:
        add_para(doc, para)

    doc.add_heading("4.4 AI Model Prediction and Confidence", level=2)
    for para in [
        "The project uses a logistic prediction model. Logistic regression is suitable for this project because it is "
        "easy to understand, fast to execute, and explainable enough for academic demonstration. The model does not "
        "require external cloud services. It loads a JSON file containing feature names, weights, and bias, then performs "
        "prediction locally.",
        "The model prediction begins with the normalized feature vector. Each feature is multiplied by its corresponding "
        "weight. The weighted values are added together with the bias. The result is passed through a sigmoid function. "
        "The sigmoid function converts any numeric input into a probability between 0 and 1. A higher value indicates "
        "that the password is more likely to be strong according to the trained feature pattern.",
        "The frontend displays AI confidence using model.predict_proba output. This is better than showing a fixed value "
        "because the confidence changes for each password. For a clearly weak password, the weak probability may be high. "
        "For a clearly strong password, the strong probability may be high. For an uncertain password, the confidence "
        "may be closer to the middle. This makes the project look and behave like an actual AI-powered analyzer.",
        "The model is trained through scripts/train_model.py using synthetic weak, medium, and strong password examples. "
        "This approach is appropriate for a student project because it avoids using real leaked password datasets, which "
        "can create privacy and legal concerns. The training examples are designed to include common weak patterns, mixed "
        "passwords, and stronger generated values. The final model file is saved in JSON format so it can be loaded by "
        "the application at runtime.",
        "The model is not the only decision maker. The analyzer combines model score with entropy score and pattern "
        "penalties. This hybrid approach is important because a small model may not perfectly understand every password "
        "case. Entropy adds mathematical strength estimation, while pattern checks add domain knowledge. The final score "
        "is therefore more reliable than any single component alone.",
    ]:
        add_para(doc, para)

    doc.add_heading("4.5 Scoring, Dynamic Colors, and User Feedback", level=2)
    for para in [
        "The scoring engine converts backend analysis into an easy-to-understand visual result. A score near 0 represents "
        "a very weak password, while a score near 100 represents an excellent password. The score is mapped to labels "
        "such as Very Weak, Weak, Fair, Strong, and Excellent. This label helps users understand the result quickly even "
        "if they do not read every technical metric.",
        "The dashboard score circle changes color according to the result. Red indicates weak passwords, orange or yellow "
        "indicates medium strength, and green indicates strong passwords. This color mapping improves usability because "
        "many users understand color faster than text. During demonstration, the evaluator can immediately see how the "
        "interface reacts when different examples are typed.",
        "The frontend also includes a loading animation while analysis is being performed. The message tells the user "
        "that the password is being analyzed. This small feature improves the feel of the application because it makes "
        "the process appear active and responsive. It also prevents confusion if a request takes a moment on a slower "
        "computer.",
        "The risk cards and suggestion cards convert technical analysis into practical advice. For example, if the "
        "password contains qwerty, the user does not only see a lower score; the system tells the user that a keyboard "
        "pattern was found. If the password contains a name, the system warns that personal information should not be "
        "included. This direct feedback supports the educational aim of the project.",
        "The entropy section shows the number of bits and gives an estimated attack-time interpretation. The estimates "
        "are approximate because attacker hardware and password hashing algorithms vary. However, the categories are "
        "useful for awareness. A low-entropy password may be described as crackable quickly, while a high-entropy "
        "password may be described as requiring years or centuries under the simplified assumptions used by the project.",
    ]:
        add_para(doc, para)

    doc.add_heading("4.6 AJAX Frontend Workflow", level=2)
    for para in [
        "The frontend uses JavaScript event listeners to detect user actions. When the password changes, the JavaScript "
        "collects the password and optional context values from the input fields. It then sends a fetch request to the "
        "backend endpoint. The request body is JSON, and the response is also JSON. This creates a clean separation "
        "between presentation and analysis logic.",
        "The AJAX design avoids a full page reload. Without AJAX, the browser would need to submit a form and refresh "
        "the page after every password check. That would make the tool feel slow and old-fashioned. With AJAX, only the "
        "dashboard content changes. The input remains in place, the page position does not jump, and the user can keep "
        "testing passwords naturally.",
        "The render function in the JavaScript file reads the response fields and updates the DOM. It sets score text, "
        "strength label, AI confidence, entropy value, attack estimates, risk messages, suggestion messages, and "
        "explanation content. This demonstrates the practical use of JSON APIs in a local project and shows how a Python "
        "model can be connected to a browser interface.",
        "The frontend also includes usability controls. The show or hide button changes the password field type so the "
        "user can inspect the typed value. The copy button uses the clipboard API to copy the generated or current "
        "password. These features are small, but they make the project feel complete because they match normal user "
        "expectations for password tools.",
    ]:
        add_para(doc, para)

    doc.add_heading("4.7 Secure Password Generator Workflow", level=2)
    for para in [
        "The generator is not a separate project; it supports the main analyzer. A password strength tool should not only "
        "tell users that a password is weak. It should also help them create a stronger password. The generator therefore "
        "produces a secure random password and immediately sends it through the same analysis pipeline.",
        "Python's secrets module is used because it is designed for security-sensitive randomness. Ordinary random "
        "functions are suitable for games, simulations, and sampling, but they are not recommended for generating "
        "passwords or tokens. Using secrets shows that the project follows secure programming principles at the code "
        "level.",
        "The generator first ensures that the password contains at least one lowercase letter, one uppercase letter, one "
        "digit, and one symbol. It then fills the remaining length using a combined alphabet and shuffles the characters "
        "securely. This avoids a predictable pattern where the first character is always lowercase or the last character "
        "is always a symbol. The final generated password is then returned to the frontend.",
        "The generator length control is bounded between 12 and 32 characters. Twelve characters is used as a practical "
        "minimum for stronger passwords in this educational project. The maximum keeps the interface readable and avoids "
        "very long values that may not fit neatly in the dashboard. The same limits are enforced in the backend so that "
        "the frontend cannot be bypassed to request an unsafe short password.",
    ]:
        add_para(doc, para)

    doc.add_heading("4.8 Sample Result Analysis", level=2)
    samples = [
        ("password123", {}),
        ("qwerty123", {}),
        ("Sachin@2026", {"name": "Sachin"}),
        ("River-Matrix-Signal-47!", {}),
        (generated, {}),
    ]
    rows = []
    for password, context in samples:
        result = analyzer.analyze(password, context=context)
        display_password = password if password != generated else "Generated password"
        rows.append(
            [
                display_password,
                str(result["score"]),
                result["label"],
                result["ai_prediction"],
                str(result["entropy_bits"]),
                result["estimated_crack_time"]["offline_fast_hash"],
            ]
        )

    add_table(
        doc,
        ["Password Tested", "Score", "Label", "AI Prediction", "Entropy Bits", "Offline Attack Estimate"],
        rows,
        "Table 6: Sample Password Analysis Results",
    )

    for para in [
        "The sample results show that the system does not treat all passwords with digits or symbols as strong. "
        "Passwords such as password123 and qwerty123 receive low scores because they contain common words, keyboard "
        "walks, and simple numeric patterns. This demonstrates that the analyzer is pattern-aware.",
        "The personal-context test Sachin@2026 is important because it shows the purpose of the optional name, email, "
        "and mobile fields. When a password contains the user's name or related personal information, the system "
        "applies a risk warning. This reflects real attack behavior because attackers often try names, phone numbers, "
        "and years when guessing passwords.",
        "The strong passphrase River-Matrix-Signal-47! receives a high score because it has good length, multiple "
        "character classes, and no obvious common keyboard walk. The generated password also receives a high score "
        "because it is produced using cryptographic randomness and includes a mix of character categories.",
        "The results support the alternative hypothesis that a feature-based AI analyzer provides more useful feedback "
        "than a simple rule-based meter. The system not only provides a score but also explains the reason behind the "
        "score. This makes the result easier to understand for normal users and more useful for educational purposes.",
        "The table also shows why password analysis must consider multiple dimensions. The example password Sachin@2026 "
        "contains uppercase, lowercase, a symbol, and digits, so a very simple composition checker might accept it. "
        "However, because it contains a personal name and a year-like value, the analyzer treats it as risky. This is "
        "one of the strongest demonstrations of the project because it shows that the optional context fields are used "
        "for meaningful analysis rather than decoration.",
        "The generated password is displayed as Generated password in the report table to avoid printing a random secret "
        "that changes every time the report builder runs. In the actual application, the generated value is visible in "
        "the interface and can be copied by the user. The report focuses on the measurable result: the generated password "
        "normally receives a high score because it is random, long enough, and mixed across character classes.",
        "The offline attack estimate in the table must be interpreted carefully. It is not a guarantee that a password "
        "will survive for exactly that time in a real attack. Real cracking speed depends on the password hash algorithm, "
        "hardware, attacker wordlists, and whether the attacker knows personal information about the user. The value is "
        "included for educational comparison because it helps users understand why small differences in entropy can make "
        "large differences in possible attack effort.",
    ]:
        add_para(doc, para)

    doc.add_heading("4.9 Screenshot Documentation", level=2)
    add_para(
        doc,
        "The screenshots included in this report document the implemented project screens. They do not include login, "
        "registration, database table, or admin-panel screenshots because those modules are not part of the current "
        "working model. Including such screenshots would misrepresent the project. Instead, the report clearly explains "
        "that the application is a stateless password analyzer and that login or database features are future scope only.",
    )

    add_picture_if_exists(doc, ASSETS / "app_dashboard.png", Inches(5.8))
    add_caption(doc, "Figure 5: Web Interface of the Analyzer")
    for para in [
        "Figure 5 shows the main dashboard of the application. The left panel contains the user input area, optional "
        "context fields, password visibility control, generator controls, and copy button. The right panel contains the "
        "analysis dashboard. This two-panel layout separates action from interpretation. The user can enter data on one "
        "side and observe results on the other side without moving to a different page.",
        "The dashboard design supports the cybersecurity theme of the project. The score circle, entropy section, attack "
        "analysis, risks, and suggestions make the project look like a security product rather than a plain classroom "
        "form. The visual structure also helps during viva because the student can point to each part of the dashboard "
        "and explain the backend logic behind it.",
    ]:
        add_para(doc, para)

    add_picture_if_exists(doc, ASSETS / "app_generated_password.png", Inches(5.8))
    add_caption(doc, "Figure 6: Generated Password Output")
    for para in [
        "Figure 6 shows the generated password workflow. When the user selects a length and requests a strong password, "
        "the frontend calls the backend generator endpoint. The generated password is placed in the password field and "
        "the analysis dashboard is updated with the generated value's score and explanation. This proves that the "
        "generator is connected to the analyzer and is not a separate static feature.",
        "The copy button improves practical usability. In a real password tool, users expect to copy a generated password "
        "without selecting it manually. The clipboard feature therefore improves the interface and demonstrates the use "
        "of browser APIs. The password is still not stored by the backend after copying; the copy operation happens on "
        "the user's device.",
    ]:
        add_para(doc, para)

    add_picture_if_exists(doc, ASSETS / "app_analysis_result.png", Inches(5.8))
    add_caption(doc, "Figure 7: Analysis Dashboard with Feedback")
    for para in [
        "Figure 7 shows the output view after analysis. The dashboard displays the strength label, score, entropy, "
        "AI prediction, AI confidence, attack estimation, detected strengths, risk warnings, improvement suggestions, "
        "and explanation. This screenshot is important because it shows the complete result expected from the working "
        "model. It also demonstrates that analysis is dynamic and does not require a page reload.",
        "The colored score circle gives immediate visual feedback. A weak password appears with a danger color, a medium "
        "password appears with a warning color, and a strong password appears with a success color. This is more effective "
        "than showing only text because the evaluator can quickly see the difference between weak and strong examples.",
    ]:
        add_para(doc, para)

    doc.add_heading("4.10 Interpretation of AI Explanation", level=2)
    for para in [
        "The AI explanation section translates model behavior into understandable points. For a weak password, it may "
        "state that the model predicted Weak with high confidence, that pattern penalties were applied for common "
        "words or keyboard patterns, and that the machine-learning score and rule-based checks agree. For a strong "
        "password, it explains that entropy and character variety are helping the password and that no major pattern "
        "risk was detected.",
        "This feature improves transparency. Users often distrust a score if they cannot understand it. By explaining "
        "confidence, entropy, and penalties, the analyzer makes the model result more educational and easier to defend "
        "during project demonstration.",
        "The explanation is also useful for debugging. If a password receives a lower score than expected, the explanation "
        "helps identify whether the cause was low entropy, a common word, repeated characters, personal information, or "
        "model uncertainty. This makes the project more maintainable because errors in scoring can be traced to specific "
        "features rather than hidden inside one large function.",
        "In a final-year project, explainability is important because the evaluator may ask how the AI component makes a "
        "decision. The answer is that the system extracts features, passes normalized values to a logistic model, receives "
        "probabilities, then combines the model result with domain rules. The displayed explanation is the user-facing "
        "form of that decision process.",
    ]:
        add_para(doc, para)

    doc.add_heading("4.11 Testing Methodology", level=2)
    for para in [
        "Testing was performed at two levels. The first level is unit testing through the Python unittest framework. Unit "
        "tests verify analyzer behavior for known weak and strong password examples, personal-context detection, repeated "
        "character detection, generator length behavior, and generator complexity. These tests help confirm that changes "
        "to the code do not break the important password-analysis rules.",
        "The second level is manual functional testing through the browser interface. Manual testing checks whether the "
        "page loads correctly, whether AJAX requests update the dashboard, whether the score circle changes color, "
        "whether the loading animation appears, whether the generated password is displayed, and whether copy-to-clipboard "
        "works. This level is necessary because some user-interface behavior cannot be fully tested by backend unit tests.",
        "Testing also covers negative and edge cases. Empty password input should not crash the application. Invalid JSON "
        "should return a structured error response. Very short generator requests should be clamped to the minimum safe "
        "length. Passwords containing repeated characters, keyboard walks, or personal details should produce relevant "
        "warnings. These cases are important because a good demo should remain stable even when the evaluator tries "
        "unusual inputs.",
        "The tests are intentionally focused on the implemented features. Since the project has no database, no login, "
        "and no admin panel, testing those modules is not applicable. The correct testing scope is the analyzer, generator, "
        "API behavior, frontend rendering, validation, and explanation output.",
    ]:
        add_para(doc, para)

    page_break(doc)
    doc.add_heading("4.12 Testing Summary", level=2)
    add_table(
        doc,
        ["Test Case", "Expected Result", "Status"],
        [
            ["Common password password123", "Low score and common-word warning", "Passed"],
            ["Strong passphrase", "High score and strong prediction", "Passed"],
            ["Personal context password", "Personal-detail warning", "Passed"],
            ["Repeated characters", "Repeated-character suggestion", "Passed"],
            ["Secure generator length 18", "Complex generated password", "Passed"],
            ["Generator short length request", "Length clamped to minimum safe length", "Passed"],
        ],
        "Table 7: Testing Summary",
    )

    doc.add_heading("4.13 Error Handling and Edge Cases", level=2)
    for para in [
        "The backend includes structured JSON error responses. If a user or tool sends invalid JSON to the API, the "
        "server returns an error response instead of crashing. If an unknown route is requested, the server returns a "
        "not-found response. This behavior is simple, but it is important for project completeness because APIs should "
        "not fail silently.",
        "The analyzer handles empty input gracefully. An empty password is not treated as a valid strong password, and "
        "the score remains zero. The interface can then guide the user to enter or generate a password. This avoids a "
        "common demo problem where empty input causes an exception or produces confusing output.",
        "The generator handles invalid length values by applying safe limits. If a very small length is requested, the "
        "backend still returns a password at least 12 characters long. This demonstrates defensive programming because "
        "the server does not depend completely on frontend validation.",
        "The frontend handles loading state and response rendering separately. This prevents partial stale results from "
        "appearing while a new request is being processed. When the response arrives, the dashboard is updated together. "
        "This makes the user experience smoother and reduces confusion during rapid testing.",
    ]:
        add_para(doc, para)
    page_break(doc)


def chapter_5(doc: Document) -> None:
    doc.add_heading("CHAPTER 5: FINDINGS AND CONCLUSION", level=1)
    doc.add_heading("5.1 Findings", level=2)
    add_para(
        doc,
        "The findings of this project are based on implementation, testing, and observation of the working password "
        "strength analyzer. The project was not designed as a survey-based research study. Instead, the main evidence "
        "comes from the software behavior: how the analyzer reacts to weak examples, strong examples, generated "
        "passwords, repeated characters, common words, keyboard patterns, entropy differences, and personal-context "
        "matches. The findings below summarize what was learned from building and testing the same application.",
    )
    add_bullets(
        doc,
        [
            "A password that satisfies uppercase, lowercase, digit, and symbol rules can still be weak if it contains a predictable pattern.",
            "Entropy is useful, but pattern detection is necessary for realistic password strength analysis.",
            "Personal-information detection improves the analyzer because many users include names, email parts, mobile numbers, or years in passwords.",
            "The logistic model provides an AI-style weak or strong prediction that can be shown as model confidence.",
            "The secure generator creates stronger passwords than most manually typed weak examples.",
            "AJAX-based dynamic analysis improves usability because the user receives feedback without reloading the page.",
            "The AI explanation section makes the model result more transparent and suitable for academic presentation.",
        ],
    )

    doc.add_heading("5.2 Detailed Discussion of Findings", level=2)
    for para in [
        "The first important finding is that simple password rules are not enough. A password may contain an uppercase "
        "letter, lowercase letters, digits, and a symbol, but it may still be weak if it is based on a name, year, or "
        "common word. Many websites still encourage users to create passwords such as Sachin@2026 because such passwords "
        "satisfy composition rules. This project shows that such rules can create a false sense of security.",
        "The second finding is that length is powerful, but only when combined with unpredictability. A longer password "
        "usually increases entropy, but repeated phrases or obvious patterns reduce practical strength. This is why the "
        "project does not reward length blindly. A long password based on passwordpassword or qwertyqwerty should not "
        "receive the same confidence as a random or carefully constructed passphrase.",
        "The third finding is that personal information is a realistic weakness. Users often include their name, mobile "
        "number, email username, birth year, college name, or simple memorable details in passwords. Attackers can often "
        "collect some of this information from social media or public profiles. By allowing optional context fields, the "
        "project demonstrates how a password checker can identify personal-information risk without storing the values.",
        "The fourth finding is that model confidence improves the project presentation. When the dashboard shows AI "
        "confidence, the output looks more like a real AI-driven application. However, confidence must come from actual "
        "model output. A static confidence percentage would be misleading. The implemented predict_proba method solves "
        "this problem by calculating probabilities from the logistic model.",
        "The fifth finding is that explainability improves trust. Users are more likely to accept feedback when the "
        "system explains why a password is weak. For example, a user may not understand why qwerty123 receives a poor "
        "score if they only look at character variety. When the analyzer points out the keyboard pattern and simple "
        "digits, the result becomes easier to understand.",
        "The sixth finding is that secure generation is a useful companion feature. Many users know that their passwords "
        "are weak but do not know how to create a better one. The generator provides an immediate alternative. Because "
        "it uses Python's secrets module, it demonstrates secure development practice and supports the overall purpose "
        "of the project.",
        "The seventh finding is that AJAX makes the application feel modern. A password-strength tool is most useful when "
        "feedback appears quickly. The dynamic interface allows a user to change the password and observe the output "
        "without page reload. This makes the project more impressive during demonstration and more comfortable for actual "
        "users.",
        "The eighth finding is that the no-database decision improves privacy. Some project reports assume that every "
        "software project must include a database, but that is not always true. For this project, not storing passwords "
        "is a security advantage. It reduces risk, keeps the scope focused, and aligns with the purpose of a local "
        "password analyzer.",
    ]:
        add_para(doc, para)

    doc.add_heading("5.3 Comparison with a Basic Password Meter", level=2)
    for para in [
        "A basic password meter usually checks only length and character variety. It may ask whether the password has "
        "uppercase letters, lowercase letters, digits, and special characters. This approach is easy to implement but "
        "does not understand user behavior. It may rate Password@123 as acceptable even though the word password and "
        "the suffix 123 are extremely common.",
        "The implemented analyzer goes beyond that approach. It calculates entropy, but it also detects common words, "
        "keyboard walks, repeated characters, year patterns, date-like values, numeric suffixes, and personal context. "
        "It then combines those signals with model prediction. This makes the result more realistic and educational. "
        "The user receives not only a score but also a reason and improvement suggestions.",
        "Another difference is the AI confidence section. A basic meter normally gives only a colored bar. This project "
        "shows model prediction and confidence, which supports the claim that the system is AI-powered. The project does "
        "not use AI as a label only; it includes feature engineering, model training, probability prediction, and "
        "explainable output.",
        "The project also includes a secure generator. Basic meters may tell users that a password is weak but leave them "
        "to solve the problem themselves. This project completes the learning loop by allowing users to generate a "
        "stronger password, copy it, and immediately see why it is stronger. That makes the application more complete "
        "for a final-year demonstration.",
    ]:
        add_para(doc, para)

    doc.add_heading("5.4 Educational Value of the Project", level=2)
    for para in [
        "The project has strong educational value because it connects multiple subjects studied in a BCA program. It "
        "uses Python programming for backend logic, web development for frontend presentation, JavaScript for dynamic "
        "interaction, JSON for data exchange, machine-learning concepts for prediction, and cybersecurity principles "
        "for password risk analysis. This combination shows practical integration rather than isolated theory.",
        "The project also teaches that security is not only about adding symbols. Many users believe that replacing a "
        "letter with a symbol automatically creates a strong password. The analyzer challenges that belief by detecting "
        "common words and predictable substitutions. This makes the project useful for awareness sessions, classroom "
        "demonstrations, and personal learning.",
        "From a programming perspective, the project demonstrates modular design. Each major responsibility is placed in "
        "a separate file. This helps maintenance and testing. If suggestions need improvement, the suggestions module "
        "can be changed without rewriting the server. If the model is retrained, the JSON file can be updated without "
        "changing the frontend. This is a good software engineering practice for a student project.",
        "From an academic perspective, the project is explainable. A large black-box AI model might be difficult to "
        "defend in viva because the student may not be able to explain why a specific result was produced. The logistic "
        "model and feature-based scoring approach make the decision process easier to present. The student can explain "
        "each feature, the entropy formula, the prediction probability, and the final score calculation.",
    ]:
        add_para(doc, para)

    doc.add_heading("5.5 Practical Usefulness", level=2)
    for para in [
        "In practical use, the application can help users understand why their passwords are weak. It is especially "
        "useful for students and non-technical users because it converts security concepts into simple messages. Instead "
        "of saying only weak, it explains whether the problem is short length, common word, keyboard pattern, repeated "
        "characters, missing variety, or personal information.",
        "The tool can also be used by a teacher or trainer to demonstrate password security. During a session, the "
        "trainer can type examples such as password123, qwerty123, a name-based password, a strong passphrase, and a "
        "generated password. The dashboard then shows how the score, entropy, AI prediction, and suggestions change. "
        "This makes the lesson visual and interactive.",
        "The application is not intended to replace enterprise password policy systems. It is a learning and analysis "
        "tool. For production systems, password strength checking should be combined with secure password hashing, "
        "multi-factor authentication, rate limiting, account lockout controls, and breached-password checking. The "
        "project report makes this boundary clear so that the project is not overstated.",
    ]:
        add_para(doc, para)

    doc.add_heading("5.6 Conclusion", level=2)
    for para in [
        "The AI-Powered Password Strength Analyzer successfully demonstrates a working software model for password "
        "security awareness. The project combines cybersecurity concepts with practical programming. It extracts "
        "password features, calculates entropy, predicts weak or strong class using a logistic model, detects weak "
        "patterns, generates intelligent suggestions, and displays the result through a professional dashboard.",
        "The system is more informative than a basic password meter because it explains risks. Instead of only saying "
        "weak or strong, it tells the user whether the password is short, contains a common word, uses a keyboard "
        "pattern, includes a year, repeats characters, or contains personal information. This supports better user "
        "learning and safer password behavior.",
        "The project also shows good privacy practice because passwords are analyzed locally and are not saved in a "
        "database or log file. The optional personal fields are used only for detection during analysis. The generator "
        "uses cryptographic randomness, which is more suitable for security than ordinary pseudo-random generation.",
        "Overall, the project meets the objective of creating a working AI-powered password analyzer suitable for BCA "
        "final-year demonstration. It connects frontend development, backend programming, model prediction, testing, "
        "and cybersecurity principles in one complete application.",
        "The most important achievement of the project is that it turns password security from a vague instruction into "
        "a measurable and explainable process. Users can see how length, character variety, entropy, patterns, personal "
        "context, and model confidence affect the final result. This makes the tool more useful than a simple checklist "
        "and supports the academic goal of demonstrating applied learning.",
        "The project also follows a responsible privacy approach. It does not store tested passwords, does not create "
        "user accounts, and does not collect unnecessary data. Optional name, email, and mobile values are used only for "
        "live analysis. This approach is appropriate because a password analyzer should minimize the amount of sensitive "
        "information it handles.",
        "The final conclusion is that the project is complete as a working model. It includes the required feature "
        "extraction functions, entropy formula, AI-style prediction, intelligent suggestions, AJAX-based frontend, "
        "dynamic colors, loading animation, copy button, secure generator, and test coverage. Future improvements can "
        "extend it, but the current implementation already satisfies the purpose of the project.",
    ]:
        add_para(doc, para)
    page_break(doc)


def chapter_6(doc: Document) -> None:
    doc.add_heading("CHAPTER 6: RECOMMENDATIONS AND LIMITATIONS OF THE STUDY", level=1)
    doc.add_heading("6.1 Recommendations", level=2)
    add_para(
        doc,
        "The recommendations in this chapter are based on the results and limitations of the implemented system. They "
        "are divided into recommendations for users, developers, institutions, and future project improvement. The "
        "recommendations are practical because the purpose of the project is not only to produce a report but also to "
        "encourage safer password behavior.",
    )
    add_bullets(
        doc,
        [
            "Users should prefer long passphrases or securely generated passwords instead of short modified dictionary words.",
            "Users should avoid names, phone numbers, email parts, dates, and years inside passwords.",
            "Applications should provide password strength feedback instead of only enforcing fixed composition rules.",
            "Organizations should combine password strength controls with multi-factor authentication.",
            "Passwords should never be stored in plaintext; production systems should use strong password hashing.",
            "Users should use a password manager so that every account can have a unique password.",
        ],
    )

    doc.add_heading("6.2 Recommendations for Users", level=2)
    for para in [
        "Users should create longer passwords or passphrases instead of short modified dictionary words. A passphrase "
        "made from multiple unrelated words is usually easier to remember and harder to guess than a short password "
        "with predictable substitutions. For example, adding a symbol to a common word is less effective than using a "
        "longer phrase that is not based on personal information.",
        "Users should avoid names, mobile numbers, email parts, birth years, college names, and other personal details. "
        "These details are often easier for attackers to collect than users assume. Social media profiles, public forms, "
        "and leaked personal data can give attackers clues. The analyzer's context-matching feature is designed to "
        "teach this lesson directly.",
        "Users should avoid reusing passwords across websites. A password that is strong on one website can become risky "
        "if the same password is used elsewhere and one of those websites is breached. Password managers are useful "
        "because they allow users to keep unique passwords for different accounts without memorizing all of them.",
        "Users should enable multi-factor authentication wherever possible. Even a strong password can be exposed through "
        "phishing or malware. Multi-factor authentication adds an additional layer of protection. The project focuses on "
        "password strength, but the report recognizes that modern authentication should use layered defenses.",
        "Users should not share passwords or store them in plain text documents. A strong password loses value if it is "
        "written in an unsafe place or sent through insecure messages. Password security includes both creating strong "
        "passwords and handling them carefully.",
    ]:
        add_para(doc, para)

    doc.add_heading("6.3 Recommendations for Developers", level=2)
    for para in [
        "Developers should avoid relying only on fixed composition rules. A rule such as one uppercase, one lowercase, "
        "one digit, and one symbol does not necessarily produce strong passwords. Developers should combine length, "
        "breached-password checking, pattern detection, and user-friendly guidance. The project demonstrates one way to "
        "combine those signals in a transparent educational tool.",
        "Developers should never store plaintext passwords. If a system actually authenticates users, passwords must be "
        "hashed using a strong password hashing algorithm with salt and suitable cost parameters. This project does not "
        "store passwords because it is not an authentication system. That design decision is also a recommendation: do "
        "not collect sensitive data unless the application truly needs it.",
        "Developers should provide helpful error messages without revealing sensitive details. In a password analyzer, "
        "the feedback should guide the user toward improvement. In a login system, however, error messages should not "
        "help attackers enumerate accounts. Context matters. The student should be able to explain this difference in "
        "viva.",
        "Developers should document assumptions behind attack-time estimates. Password cracking speed depends on many "
        "factors, including hardware, hash type, rate limits, and whether the attacker is online or offline. A project "
        "should not present estimates as exact guarantees. The implemented analyzer treats them as educational categories.",
        "Developers should keep security tools explainable. If users cannot understand why a password is weak, they may "
        "ignore the warning. The suggestions and explanation modules in this project show how technical analysis can be "
        "converted into understandable feedback.",
    ]:
        add_para(doc, para)

    doc.add_heading("6.4 Recommendations for Institutions", level=2)
    for para in [
        "Educational institutions can use password analyzers as part of cyber-awareness training. Many students use "
        "simple passwords for email, learning portals, cloud storage, and social media. A visual tool can show the risk "
        "more effectively than a lecture alone. The project can therefore be demonstrated in classrooms or workshops.",
        "Institutions should encourage password managers and multi-factor authentication for student and staff accounts. "
        "Password strength is important, but it should be part of a broader security culture. Training should also cover "
        "phishing, safe device use, account recovery security, and the danger of password reuse.",
        "Institutions should avoid collecting actual student passwords for demonstrations. This project is appropriate "
        "because it can run locally and does not save tested passwords. If used in a class, students should test sample "
        "passwords or newly generated examples rather than their real account passwords.",
    ]:
        add_para(doc, para)

    doc.add_heading("6.5 Limitations", level=2)
    add_bullets(
        doc,
        [
            "The model is trained on synthetic examples and does not include real leaked-password datasets because such datasets may contain sensitive information.",
            "The crack-time estimates are approximate because attacker hardware and hashing algorithms vary.",
            "The system does not connect to an online breached-password database.",
            "The application is an educational local tool, not a complete enterprise identity security system.",
            "The analyzer uses a lightweight logistic model for transparency rather than a large deep-learning model.",
        ],
    )

    doc.add_heading("6.6 Detailed Limitations", level=2)
    for para in [
        "The first limitation is the training data. The project uses synthetic password examples rather than real leaked "
        "password datasets. This is a responsible choice for a student project because leaked password lists can contain "
        "real user secrets and may create ethical or legal concerns. However, it also means the model may not learn the "
        "full diversity of real-world password habits.",
        "The second limitation is that the model is lightweight. Logistic regression is transparent and easy to explain, "
        "but it cannot capture every complex password pattern. A larger model or a specialized password-strength library "
        "could detect additional patterns. The project balances accuracy with explainability and ease of demonstration.",
        "The third limitation is that crack-time estimates are approximate. The same password can be much harder or "
        "easier to crack depending on the environment. Online attacks may be slowed by rate limits and account lockouts. "
        "Offline attacks against fast hashes may be much faster. Therefore, the project presents crack time as a guide, "
        "not as a precise prediction.",
        "The fourth limitation is that the application does not connect to a breached-password database. A password may "
        "look strong based on length and symbols but still be unsafe if it has appeared in a breach. Adding privacy-preserving "
        "breach checking would improve the tool, but it would require an external service and careful privacy design.",
        "The fifth limitation is that the project is local and educational. It is not deployed as a hardened production "
        "service. Production deployment would require HTTPS, secure headers, logging controls, rate limiting, deployment "
        "configuration, monitoring, and a stronger threat model. These topics are outside the current project scope.",
        "The sixth limitation is that the interface is built for a single-user local workflow. It does not include user "
        "roles, admin panels, report history, or analytics. This is a limitation only if the project is treated as an "
        "organizational platform. For the current objective, the simpler design is appropriate and safer.",
    ]:
        add_para(doc, para)

    doc.add_heading("6.7 Future Scope", level=2)
    add_bullets(
        doc,
        [
            "Add privacy-preserving breached-password checking through a k-anonymity API.",
            "Add browser-only analysis mode so that even local API transmission is avoided.",
            "Add multilingual dictionary detection for Indian languages and common regional names.",
            "Add admin-configurable password policies for organizations.",
            "Add PDF export of password awareness reports for training sessions.",
            "Improve model training with larger safe datasets and comparative evaluation.",
        ],
    )

    doc.add_heading("6.8 Detailed Future Enhancements", level=2)
    for para in [
        "A future version can add privacy-preserving breached-password checking. The safest approach would be to use a "
        "k-anonymity method where only a hash prefix is sent to an external service. The full password should never be "
        "sent to any server. This enhancement would allow the analyzer to warn users when a password has already appeared "
        "in known breaches.",
        "A future version can add a browser-only analysis mode. In that design, feature extraction and scoring would run "
        "completely in JavaScript or WebAssembly. This would remove even local backend transmission. It would be useful "
        "for users who are very cautious about typing passwords into any server process, even one running on localhost.",
        "A future version can improve the model with a larger safe dataset. The project can generate more synthetic "
        "examples, include more passphrase structures, include regional naming patterns without storing real personal "
        "data, and compare logistic regression with decision trees or gradient boosting. The report can then include "
        "accuracy, precision, recall, and confusion matrix evaluation.",
        "A future version can add multilingual and regional pattern detection. Many users in India include names, local "
        "words, city names, college names, or language-specific terms in passwords. A safe dictionary of common terms "
        "could improve risk detection while still avoiding the storage of individual personal data.",
        "A future version can add an optional organizational policy mode. In that mode, an admin could define minimum "
        "length, required categories, banned words, and minimum score. However, such a feature should still avoid storing "
        "real passwords. It should only store policy configuration and anonymous aggregate statistics if analytics are "
        "needed.",
        "A future version can export awareness reports. For example, after a training session, the tool could export a "
        "PDF explaining common mistakes and recommended practices. The exported report should not include actual tested "
        "passwords. It should use sample values or anonymized categories only.",
        "A future version can improve accessibility. Keyboard navigation, clearer focus indicators, screen-reader labels, "
        "and high-contrast mode would make the tool easier to use for more people. Since password security is relevant "
        "to all users, accessibility should be considered part of quality.",
        "A future version can add stronger automated frontend tests. Browser automation could verify that the dashboard "
        "updates correctly, that the spinner appears, that the score circle changes color, and that the copy button works. "
        "This would make future UI changes safer.",
    ]:
        add_para(doc, para)

    doc.add_heading("6.9 Final Recommendation", level=2)
    for para in [
        "The final recommendation is to submit the project as a focused AI-powered password strength analyzer, not as a "
        "general login or database management system. The strongest part of the project is the working model and the "
        "clear dashboard. Adding unrelated features only for appearance could reduce quality and create factual problems "
        "in the report. The documentation should therefore emphasize the real implemented features: feature extraction, "
        "entropy, model probability, intelligent suggestions, secure generation, AJAX, dynamic UI feedback, and privacy.",
        "For the final college submission, the student should attach the plagiarism report generated by the approved "
        "tool, update page numbers in the table of contents if required by the guide, add the guide's signature on the "
        "certificate page, and insert any college-specific enrollment or program details that are still blank. The "
        "technical content is now expanded enough to support a longer report, but administrative fields must still be "
        "completed by the student before physical or portal submission.",
    ]:
        add_para(doc, para)
    page_break(doc)


def add_references(doc: Document) -> None:
    doc.add_heading("BIBLIOGRAPHY / REFERENCES", level=1)
    refs = [
        "National Institute of Standards and Technology. (2025). Digital Identity Guidelines: Authentication and Authenticator Management (NIST SP 800-63B-4). https://pages.nist.gov/800-63-4/sp800-63b.html",
        "OWASP Foundation. (2026). Authentication Cheat Sheet. https://cheatsheetseries.owasp.org/cheatsheets/Authentication_Cheat_Sheet.html",
        "OWASP Foundation. (2026). Password Storage Cheat Sheet. https://cheatsheetseries.owasp.org/cheatsheets/Password_Storage_Cheat_Sheet.html",
        "Wheeler, D. L. (2016). zxcvbn: Low-Budget Password Strength Estimation. 25th USENIX Security Symposium, 157-173. https://www.usenix.org/conference/usenixsecurity16/technical-sessions/presentation/wheeler",
        "Python Software Foundation. (2026). secrets - Generate secure random numbers for managing secrets. https://docs.python.org/3/library/secrets.html",
    ]
    for ref in refs:
        add_para(doc, ref)
    page_break(doc)


def add_appendices(doc: Document) -> None:
    doc.add_heading("APPENDIX", level=1)
    doc.add_heading("Appendix A: How to Run the Project", level=2)
    add_numbered(
        doc,
        [
            r"Open PowerShell in C:\Users\Sachin Nishad\Documents\New project.",
            "Run the command: python app.py",
            "Open http://127.0.0.1:8000 in the browser.",
            "Enter a password or click Generate Strong Password.",
            "Observe score, AI model confidence, entropy, risks, suggestions, and AI explanation.",
        ],
    )

    doc.add_heading("Appendix B: Important Feature Extraction Code", level=2)
    code_lines = [
        "length = len(password)",
        "lower_count = sum(ch.islower() for ch in password)",
        "upper_count = sum(ch.isupper() for ch in password)",
        "digit_count = sum(ch.isdigit() for ch in password)",
        "symbol_count = sum(not ch.isalnum() for ch in password)",
        "charset_size = _charset_size(password)",
        "entropy_bits = length * math.log2(charset_size) if password else 0.0",
        "repeat_ratio = (length - unique_chars) / length if length else 0.0",
    ]
    for line in code_lines:
        p = doc.add_paragraph(style="Code Block")
        p.add_run(line)

    doc.add_heading("Appendix C: Project Folder Structure", level=2)
    add_table(
        doc,
        ["File / Folder", "Purpose"],
        [
            ["app.py", "Runs the local backend server and API endpoints."],
            ["password_ai/feature_engineering.py", "Extracts numeric password features."],
            ["password_ai/model.py", "Loads and executes the logistic prediction model."],
            ["password_ai/analyzer.py", "Combines model, entropy, risks, suggestions, and explanations."],
            ["password_ai/generator.py", "Creates secure random passwords."],
            ["static/", "Contains HTML, CSS, and JavaScript frontend files."],
            ["tests/", "Contains unit tests for analyzer and generator."],
        ],
    )

    doc.add_heading("Appendix D: API Request and Response Documentation", level=2)
    for para in [
        "The application communicates through simple JSON APIs. The browser sends a POST request to the backend and the "
        "backend returns a structured JSON response. This appendix documents the important fields so that the project "
        "can be understood and tested without reading the entire source code. The API is local and intended for the "
        "same-machine browser interface at http://127.0.0.1:8000.",
        "The /api/analyze endpoint accepts a password and optional personal context. The context values may include name, "
        "email, and mobile number. These values are used only during the current analysis request. They are not saved to "
        "a database and are not written to a report. The purpose of the context object is to detect risky personal "
        "information inside the password.",
    ]:
        add_para(doc, para)
    for line in [
        "{",
        '  "password": "Sachin@2026",',
        '  "context": {',
        '    "name": "Sachin",',
        '    "email": "sachin@example.com",',
        '    "mobile": "9876543210"',
        "  }",
        "}",
    ]:
        p = doc.add_paragraph(style="Code Block")
        p.add_run(line)
    for para in [
        "The response contains the calculated score, label, entropy, AI prediction, confidence, estimated attack time, "
        "strength messages, risk messages, suggestions, and explanation. The frontend does not calculate these values "
        "again. It only displays the response. This keeps the password-strength logic in the backend and makes the "
        "frontend easier to maintain.",
        "The /api/generate endpoint accepts a desired length. The backend clamps the length to a safe range, creates a "
        "secure password using the generator module, analyzes that password, and returns both the generated value and "
        "the analysis object. This allows the frontend to display the generated password and its score in one workflow.",
    ]:
        add_para(doc, para)
    for line in [
        "{",
        '  "length": 18',
        "}",
    ]:
        p = doc.add_paragraph(style="Code Block")
        p.add_run(line)
    add_para(
        doc,
        "A typical analysis response includes fields similar to score, label, ai_prediction, ai_confidence, entropy_bits, "
        "estimated_crack_time, strengths, risks, suggestions, and explanation. The exact values depend on the tested "
        "password. This structure is suitable for AJAX because each field can be rendered into a specific dashboard "
        "section without refreshing the page.",
    )

    doc.add_heading("Appendix E: Module-Level Code Documentation", level=2)
    for para in [
        "The app.py file creates the server class and routes requests. The do_GET method serves the main HTML page and "
        "static assets. The do_POST method handles API routes. Helper methods read JSON, write JSON responses, and send "
        "error objects. The file imports PasswordStrengthAnalyzer and generate_secure_password from the password_ai "
        "package so the web server can use the analysis model.",
        "The feature_engineering.py module should be studied carefully because it represents the feature extraction part "
        "of the project. Important functions calculate character classes, charset size, entropy, repeated runs, sequence "
        "patterns, common words, keyboard walks, suffix digits, and personal-context matches. These features are later "
        "used by the model and by the suggestion engine.",
        "The model.py module is intentionally small. It defines the logistic scoring operation, loads model data from a "
        "JSON file, and exposes probability output. This makes the AI component easy to explain. The model does not call "
        "an external API and does not require internet access. The project can therefore be demonstrated offline after "
        "the code and dependencies are available.",
        "The analyzer.py module is the decision coordinator. It combines the model result with entropy and pattern rules. "
        "It also formats the result for the frontend. This file is where score, label, AI confidence, estimated crack "
        "time, strengths, risks, and explanation come together. If the project is extended in the future, this module "
        "will likely be the main place for changing scoring policy.",
        "The suggestions.py module is responsible for improvement advice. A good suggestion should be specific, short, "
        "and directly connected to the detected weakness. For example, if the analyzer detects a common word, the "
        "suggestion should tell the user to avoid dictionary words. If the analyzer detects personal context, the "
        "suggestion should tell the user to remove names, email parts, phone numbers, or years.",
        "The generator.py module uses secure randomness. The most important point is that it uses secrets rather than "
        "ordinary random. It also ensures character variety and applies safe length bounds. The generated password is "
        "not stored. It is returned to the frontend and then analyzed like any other password.",
        "The frontend JavaScript file connects the interface to the backend. It collects values from input fields, sends "
        "fetch requests, handles loading state, receives JSON, and updates the page. It also supports show or hide "
        "password, copy-to-clipboard, generator control, dynamic score color, and dashboard rendering. This file proves "
        "that the backend model is connected to an interactive UI.",
        "The CSS file gives the application a professional dashboard appearance. It controls the two-panel layout, score "
        "circle, metric cards, buttons, suggestions, colors, spacing, and responsive behavior. Although CSS is not the "
        "AI part of the project, it is important for presentation because a good final-year project should be easy to "
        "understand visually.",
    ]:
        add_para(doc, para)

    doc.add_heading("Appendix F: Viva Voce Preparation", level=2)
    for para in [
        "Question: Why is the project called AI-powered? Answer: The project is AI-powered because it uses feature "
        "engineering and a trained logistic model to predict whether a password is weak or strong. The model returns "
        "probabilities through predict_proba, and the application displays AI prediction and confidence. The system also "
        "uses rule-based domain checks for better explanation, so it is a hybrid AI and cybersecurity project.",
        "Question: What is the entropy formula used in the project? Answer: The entropy formula is Entropy = L x log2(N). "
        "Here L is the password length and N is the estimated character set size. If a password uses more character "
        "classes, N increases. If the password is longer, L increases. Higher entropy usually means a larger theoretical "
        "search space.",
        "Question: Why is entropy alone not enough? Answer: Entropy assumes random selection, but most users create "
        "passwords using patterns. A password may have symbols and digits but still contain a common word, name, date, "
        "or keyboard walk. The project therefore combines entropy with pattern detection, model prediction, and targeted "
        "suggestions.",
        "Question: Why does the project not use a database? Answer: The project is a password analyzer, not an account "
        "management system. It does not need to store users, passwords, or reports. Not using a database is safer because "
        "tested passwords and personal context should not be stored unnecessarily.",
        "Question: What is the role of AJAX in the project? Answer: AJAX allows the frontend to send password-analysis "
        "requests to the backend and update the dashboard without reloading the page. This makes the application feel "
        "live and interactive. The Fetch API sends JSON data and receives JSON analysis results.",
        "Question: How are intelligent suggestions created? Answer: The suggestion module checks the extracted features. "
        "If the password is short, it recommends increasing length. If a character class is missing, it recommends adding "
        "variety. If common words, personal details, repeated characters, dates, or keyboard walks are found, it gives "
        "specific advice related to that weakness.",
        "Question: Why is Python's secrets module used? Answer: The secrets module is designed for secure random values. "
        "Password generation is security-sensitive, so ordinary random functions should not be used. The generator uses "
        "secrets to select and shuffle password characters.",
        "Question: What are the main limitations? Answer: The model is trained on synthetic examples, crack-time estimates "
        "are approximate, breached-password checking is not included, and the project is a local educational tool rather "
        "than a production identity platform. These limitations are clearly documented in the report.",
        "Question: How can the project be improved in the future? Answer: Future improvements can include privacy-preserving "
        "breached-password checking, browser-only analysis, multilingual pattern detection, larger safe training data, "
        "frontend automation tests, accessibility improvements, and optional anonymous awareness reports.",
    ]:
        add_para(doc, para)

    doc.add_heading("Appendix G: Final Submission Checklist and Originality Note", level=2)
    for para in [
        "Before submission, the student should review the report with the project guide and confirm that all administrative "
        "details are complete. The title page should contain the correct student name, enrollment number, program name, "
        "academic session, guide name, and submission date. The declaration and certificate pages should be signed where "
        "required. If the university portal asks for a separate plagiarism report, that report should be generated using "
        "the approved tool and attached with the final file.",
        "The table of contents page should be checked after final pagination. Word processors can change page numbers "
        "when screenshots, spacing, or signatures are added. If automatic fields are not used, the page numbers should "
        "be manually updated before printing or uploading. The list of tables and list of figures should also be checked "
        "so that table captions and figure captions match the body of the report.",
        "Screenshots should be taken from the actual running project, not from unrelated websites or templates. The "
        "included screenshots must show the implemented dashboard, generated password workflow, and analysis output. "
        "Screenshots of login, register, database tables, or admin panels should not be added unless those features are "
        "actually implemented. This protects the project from mismatch during viva.",
        "For originality, the report uses project-specific wording and explains the actual source code created for this "
        "application. Technical terms such as entropy, logistic model, AJAX, and password hashing are standard concepts "
        "and may appear in many cybersecurity documents, but the explanations here are written in relation to this "
        "specific project. Direct copying from websites should be avoided. External sources should remain limited to "
        "the reference section and should be cited properly.",
        "The codebase should be kept with the report during demonstration. The student should be ready to open app.py, "
        "feature_engineering.py, model.py, analyzer.py, suggestions.py, generator.py, and static/app.js. During viva, it "
        "is better to explain a small piece of real working code clearly than to claim features that are not present. "
        "The strength of this project is that the model, frontend, suggestions, generator, and tests are connected.",
        "A final dry run should be performed before submission. Start the project with python app.py, open the browser "
        "at http://127.0.0.1:8000, test one weak password, test one personal-information password, generate a strong "
        "password, copy the generated password, and run the unit tests. If these steps work, the project demonstration "
        "will match the documentation in this report.",
        "The final report should be treated as a living academic document until submission day. If the guide suggests "
        "small wording, formatting, or screenshot changes, those updates should be made in the same report without "
        "changing the project topic or inventing unsupported modules.",
    ]:
        add_para(doc, para)


def add_picture_if_exists(doc: Document, path: Path, width) -> None:
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=width)
    else:
        add_para(doc, f"[Screenshot placeholder: {path.name}]")


if __name__ == "__main__":
    main()
